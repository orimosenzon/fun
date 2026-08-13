"""טסטים לשני השינויים ב-lib/haskala_grading מ-11/8/2026.

שניהם באו מאבישי:
  • ספירת מילים — רק כתב יד, בלי טקסט מודפס מדף הבחינה
  • טקסט שצריך להימחק — בסוגריים, קו חוצה והדגשה צהובה

בלי רשת ובלי קריאות למודל.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from haskala_grading import core, report  # noqa: E402


# ─── word count: handwriting only ───────────────────────────────────────────

def _page(*lines):
    """lines are (text, printed) pairs."""
    return [{"page": 1, "lines": [{"text": t, "printed": p} for t, p in lines]}]


def test_printed_lines_are_not_counted():
    """The exam's own question used to count as the student's writing."""
    pages = _page(("Write about your summer vacation in 70-90 words", True),
                  ("Last summer I went to Greece with my family", False))
    assert core.count_words(pages) == 9


def test_a_page_that_is_all_printed_counts_zero():
    pages = _page(("Ministry of Education", True), ("Module C 016382", True))
    assert core.count_words(pages) == 0


def test_lines_without_the_flag_are_still_counted():
    """Back-compat: results stored before the flag existed, and callers that
    build pages by hand, must not lose their text to a field they never set."""
    pages = [{"page": 1, "lines": [{"text": "one two three"}]}]
    assert core.count_words(pages) == 3


def test_the_deduction_this_was_really_about():
    """A sixty-word answer under a 70–90 rule earns a deduction. Forty words of
    printed prompt used to push it to a hundred and cancel that."""
    student = " ".join(["word"] * 60)
    printed = " ".join(["instructions"] * 40)
    assert core.count_words(_page((printed, True), (student, False))) == 60
    # and the old behaviour, for contrast
    assert core.count_words(_page((printed, False), (student, False))) == 100


def test_the_ocr_prompt_asks_for_the_flag():
    """The count is only honest while the OCR is actually classifying lines."""
    prompt = core.build_ocr_prompt("en") if hasattr(core, "build_ocr_prompt") else None
    if prompt is None:                      # prompt builder is named differently
        import re
        src = open(core.__file__, encoding="utf-8").read()
        assert '"printed"' in src and "מודפסת" in src
    else:
        assert "printed" in prompt


# ─── deletions: bracketed, struck through, yellow ───────────────────────────

def _evaluation(delete: bool):
    return {"criteria": [{
        "name": "Language Use", "color": "3366CC", "score": 5, "max_score": 10,
        "level": "fair", "feedback": "", "feedback_secondary": "",
        "issues": [{"line_ref": "p1-l0", "quote": "very very",
                    "comment": "repeated word", "delete": delete}],
    }]}


_PAGES = [{"page": 1, "lines": [{"text": "It was very very good", "printed": False}]}]


def test_delete_flag_reaches_the_span():
    spans = report.resolve_issue_spans(_PAGES, _evaluation(True))
    assert spans[(1, 0)][0]["delete"] is True
    spans = report.resolve_issue_spans(_PAGES, _evaluation(False))
    assert spans[(1, 0)][0]["delete"] is False


def _render(delete: bool):
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    spans = report.resolve_issue_spans(_PAGES, _evaluation(delete))[(1, 0)]
    report._fill_paragraph_with_spans(p, _PAGES[0]["lines"][0]["text"], spans,
                                      include_notes=False)
    return p


def _shading(run):
    from docx.oxml.ns import qn
    rPr = run._r.find(qn("w:rPr"))
    shd = rPr.find(qn("w:shd")) if rPr is not None else None
    return shd.get(qn("w:fill")) if shd is not None else None


def test_deleted_text_is_bracketed_struck_and_yellow():
    marked = [r for r in _render(True).runs if r.text.startswith("[")]
    assert len(marked) == 1
    run = marked[0]
    assert run.text == "[very very]"
    assert run.font.strike is True
    assert _shading(run) == report.DELETE_HIGHLIGHT


def test_a_normal_issue_is_untouched_by_any_of_that():
    """The three marks mean "remove this". A correction that only needs
    rewriting must not be shown to a student as text to delete."""
    runs = _render(False).runs
    marked = [r for r in runs if r.text == "very very"]
    assert len(marked) == 1
    assert not marked[0].font.strike
    assert _shading(marked[0]) != report.DELETE_HIGHLIGHT
    assert not any(r.text.startswith("[") for r in runs)


def test_yellow_replaces_the_criterion_colour_rather_than_mixing():
    """Avishai's call: two shades over one span is unreadable, and the
    criterion is still named in the note that follows the span."""
    crit_tint = report._tint_hex("3366CC", mix=0.35)
    run = [r for r in _render(True).runs if r.text.startswith("[")][0]
    assert _shading(run) == report.DELETE_HIGHLIGHT != crit_tint


# ─── erasures: not counted, shown as a short label ──────────────────────────
# Avishai's 2026-08-12 submission: the count read 143 against a composition of
# 127, and eleven of the difference were words the student had struck through.

def _erased_page(*lines):
    """lines are (text, erased_fragments) pairs, all handwriting."""
    return [{"page": 1, "lines": [{"text": t, "printed": False, "erased": list(e)}
                                  for t, e in lines]}]


def test_crossed_out_words_are_not_counted():
    pages = _erased_page(
        ("important language, Because the English", []),
        ("la internashunal language and all the contre",
         ["la internashunal language and all the contre"]),
        ("a international language and all the", []),
    )
    assert core.count_words(pages) == 11   # 5 + 0 + 6, not 5 + 7 + 6


def test_a_fragment_in_the_middle_of_a_line_leaves_its_neighbours_apart():
    """Cutting the fragment out rather than blanking it would fuse the words
    on either side into one and quietly undercount by one."""
    line = {"text": "if you want to, Trave Travel or learn", "erased": ["Trave"]}
    assert core.strip_erased(line).split() == \
        ["if", "you", "want", "to,", "Travel", "or", "learn"]
    assert core.count_words([{"page": 1, "lines": [line]}]) == 7


def test_a_fragment_the_model_did_not_copy_verbatim_is_ignored():
    """The OCR sometimes returns its own reading of the scribble ('anywage')
    instead of the substring it transcribed ('anyway'). Unlocatable, so it
    cannot be excluded — and must not silently subtract a word elsewhere."""
    line = {"text": "she help you anyway any way", "erased": ["anywage"]}
    assert core.erased_fragments(line) == []
    assert core.count_words([{"page": 1, "lines": [line]}]) == 6


def test_pages_without_the_field_are_unaffected():
    pages = [{"page": 1, "lines": [{"text": "one two three", "printed": False}]}]
    assert core.count_words(pages) == 3


def test_the_ocr_prompt_asks_for_the_erasures():
    """The exclusion only works while the OCR is actually reporting them."""
    assert "erased" in core.build_ocr_prompt("en")


def test_the_evaluator_is_told_to_ignore_what_the_student_cancelled():
    """Otherwise the student is marked down for a draft they withdrew — which
    is exactly what happened on 2026-08-12 ('garbled repeated attempt')."""
    pages = _erased_page(("good text and bad text", ["and bad text"]))
    transcript = core.pages_to_plain_text(pages)
    assert "⟨מחוק: and bad text⟩" in transcript
    assert "מחוק" in core.build_eval_prompt("he", "en")


def _erased_render(lang="en"):
    from docx import Document
    pages = _erased_page(("First of all the contre a international", ["the contre"]))
    spans = report.merge_erased_spans({}, report.resolve_erased_spans(pages))[(1, 0)]
    doc = Document()
    p = doc.add_paragraph()
    report._fill_paragraph_with_spans(
        p, pages[0]["lines"][0]["text"], spans,
        erased_label=report.DOCX_LABELS[lang]["erased_text"])
    return p


def test_erased_text_is_replaced_by_the_short_label():
    """Ori/Avishai (2026-08-13): the label, struck and yellow. Not the garbled
    fragment itself, and not a sentence explaining what it was."""
    runs = _erased_render().runs
    marked = [r for r in runs if r.font.strike]
    assert len(marked) == 1
    assert marked[0].text == "erased text"
    assert _shading(marked[0]) == report.DELETE_HIGHLIGHT
    assert "the contre" not in "".join(r.text for r in runs)


def test_the_label_carries_no_note():
    """`include_notes` is on by default — an erasure still gets no explanation
    after it, unlike every other kind of marked span."""
    assert "(" not in "".join(r.text for r in _erased_render().runs)


def test_the_label_follows_the_report_language():
    assert [r for r in _erased_render("he").runs if r.font.strike][0].text == "טקסט מחוק"


def test_an_issue_quoting_erased_text_is_dropped():
    """A correction to writing the student already withdrew tells them to fix
    something that is not part of their answer."""
    pages = _erased_page(("she help you anyway", ["anyway"]))
    issues = {(1, 0): [{"start": 13, "end": 19, "color": "3366CC",
                        "comment": "spelling", "criterion": "Mechanics",
                        "delete": False}]}
    merged = report.merge_erased_spans(issues, report.resolve_erased_spans(pages))
    assert [s.get("erased") for s in merged[(1, 0)]] == [True]


def test_an_issue_elsewhere_on_the_line_survives():
    pages = _erased_page(("she help you anyway", ["anyway"]))
    issues = {(1, 0): [{"start": 4, "end": 8, "color": "3366CC",
                        "comment": "verb form", "criterion": "Language Use",
                        "delete": False}]}
    merged = report.merge_erased_spans(issues, report.resolve_erased_spans(pages))
    assert len(merged[(1, 0)]) == 2
    assert [s["start"] for s in merged[(1, 0)]] == [4, 13]


if __name__ == "__main__":
    failures = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            try:
                fn()
                print(f"  ok   {name}")
            except Exception as e:
                failures += 1
                print(f"  FAIL {name}: {type(e).__name__}: {e}")
    print(f"\n{failures} failure(s)")
    sys.exit(1 if failures else 0)
