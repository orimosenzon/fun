"""math-checker — בדיקת תרגילי מתמטיקה/גאומטריה בכתב יד (מוצר #2, חברת השכלה).

בניגוד לבודק-השפה (products/checker), כאן אין חיתוך שורה-שורה. תרגיל מתמטי
הוא יחידה הוליסטית — סרטוטים, הוכחות גאומטריות ומעברים אלגבריים פרושים ב-2D.
לכן העמוד המלא נשלח ל-Opus שמחזיר הבנה מובנית (JSON): רשימת תרגילים, ולכל
אחד את צעדי הפתרון ב-LaTeX, הערכת תקינות לכל צעד, ומשוב.
"""
from __future__ import annotations

import base64
import datetime
import hmac
import io
import json
import logging
import os
import queue
import threading
import time
import uuid

import anthropic
import fitz  # PyMuPDF
from dotenv import load_dotenv
from flask import Flask, Response, jsonify, render_template, request
from PIL import Image, ImageOps

load_dotenv(override=True)

LOG_PATH = os.path.join(os.path.dirname(__file__), "math-checker.log")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(message)s",
    handlers=[logging.FileHandler(LOG_PATH, mode="a"), logging.StreamHandler()],
)
log = logging.getLogger("math-checker")

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024

DESKTOP_MODE = os.environ.get("MATH_CHECKER_DESKTOP") == "1"
BASIC_USER = os.environ.get("HASKALA_USER")
BASIC_PASS = os.environ.get("HASKALA_PASS")
if not DESKTOP_MODE and not (BASIC_USER and BASIC_PASS):
    log.warning(
        "running without Basic Auth — set HASKALA_USER/HASKALA_PASS "
        "before exposing this publicly"
    )


@app.before_request
def _require_basic_auth():
    if DESKTOP_MODE or not (BASIC_USER and BASIC_PASS):
        return
    if request.path.startswith("/static/"):
        return
    auth = request.authorization
    if (
        auth
        and auth.type == "basic"
        and hmac.compare_digest(auth.username or "", BASIC_USER)
        and hmac.compare_digest(auth.password or "", BASIC_PASS)
    ):
        return
    return Response(
        "נדרשת הזדהות",
        401,
        {"WWW-Authenticate": 'Basic realm="math-checker"'},
    )


client = anthropic.Anthropic()

RENDER_DPI = 200          # A4 @ 200 DPI → ~1654×2339 px (well under Opus vision limit)
MAX_EDGE = 2200           # longest edge sent to Opus for analysis
ORIENT_MAX_EDGE = 900     # longest edge sent to Haiku for fast orientation check
PREVIEW_MAX_EDGE = 1400   # longest edge stored in result for display in browser

# Analysis models the teacher can pick in the UI. Holistic math/geometry
# grading is a hard reasoning task (proofs, 2D diagrams, multi-step algebra),
# so only strong reasoners do it well. The full provider line-up from
# products/checker is exposed for the current testing phase, but the cheap /
# mini / free tiers are flagged "(לא מומלץ)" — empirically they miss real
# calculation/logic errors on math pages. Opus is the quality default.
#
# Each UI key maps to (provider, model_id) so analyze_page can dispatch across
# Anthropic / Google / Groq / Azure (mirrors products/checker's abstraction).
MODELS = {
    "opus":             "Claude Opus 4.8 · איכות מרבית",
    "sonnet":           "Claude Sonnet 4.6 · מהיר וזול ~5×",
    "gemini-flash":     "Gemini 2.5 Flash · מהיר וזול",
    "haiku":            "Claude Haiku 4.5 · הכי זול (לא מומלץ)",
    "gemini-lite":      "Gemini 2.5 Flash-Lite · זול מאוד (לא מומלץ)",
    "groq-scout":       "Groq Llama 4 Scout · חינמי (לא מומלץ)",
    "azure-gpt41-mini": "GPT-4.1-mini · Azure (לא מומלץ)",
}

# provider kind + concrete model id per UI key
_ANTHROPIC_IDS = {
    "opus":   "claude-opus-4-8",
    "sonnet": "claude-sonnet-4-6",
    "haiku":  "claude-haiku-4-5-20251001",
}
_GEMINI_IDS = {
    "gemini-flash": "gemini-2.5-flash",
    "gemini-lite":  "gemini-2.5-flash-lite",
}
_GROQ_IDS = {
    "groq-scout": "meta-llama/llama-4-scout-17b-16e-instruct",
}
_AZURE_KEYS = {"azure-gpt41-mini"}

DEFAULT_MODEL = "opus"
ORIENT_MODEL  = "claude-haiku-4-5-20251001"  # orientation check is always Haiku (cheap)


def resolve_model(key: str | None) -> tuple[str, str]:
    """Map a UI model key to (canonical_key, hebrew_label). Falls back to default."""
    k = key if key in MODELS else DEFAULT_MODEL
    return k, MODELS[k]

ACCEPTED_EXTS = (".pdf", ".jpg", ".jpeg", ".png")

JOB_TTL = 3600  # seconds — results stay in memory for result refreshes
JOBS: dict[str, dict] = {}


# ─── file handling ────────────────────────────────────────────────────────────

def pdf_to_pages(pdf_bytes: bytes) -> list[Image.Image]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    imgs = []
    for page in doc:
        pix = page.get_pixmap(dpi=RENDER_DPI)
        imgs.append(Image.frombytes("RGB", (pix.width, pix.height), pix.samples))
    doc.close()
    return imgs


def file_to_pages(data: bytes, ext: str) -> list[Image.Image]:
    if ext == ".pdf":
        return pdf_to_pages(data)
    img = ImageOps.exif_transpose(Image.open(io.BytesIO(data)))
    return [img.convert("RGB")]


def _downscale(img: Image.Image, max_edge: int) -> Image.Image:
    long_edge = max(img.size)
    if long_edge <= max_edge:
        return img
    scale = max_edge / long_edge
    return img.resize(
        (round(img.width * scale), round(img.height * scale)),
        Image.LANCZOS,
    )


def _jpeg_bytes(img: Image.Image, max_edge: int | None = None, quality: int = 90) -> bytes:
    if max_edge:
        img = _downscale(img, max_edge)
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=quality)
    return buf.getvalue()


def _img_to_b64_jpeg(img: Image.Image, max_edge: int | None = None, quality: int = 90) -> str:
    return base64.standard_b64encode(_jpeg_bytes(img, max_edge, quality)).decode()


# ─── optional non-Claude providers (lazy: a missing key only bites if picked) ──
_gemini_client = None
_groq_client = None
_azure_client = None


def _gemini():
    global _gemini_client
    if _gemini_client is None:
        from google import genai
        _gemini_client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    return _gemini_client


def _groq():
    global _groq_client
    if _groq_client is None:
        from groq import Groq
        _groq_client = Groq(api_key=os.environ["GROQ_API_KEY"])
    return _groq_client


def _azure():
    global _azure_client
    if _azure_client is None:
        from openai import OpenAI
        base = os.environ["AZURE_OPENAI_ENDPOINT"].rstrip("/") + "/"
        _azure_client = OpenAI(base_url=base, api_key=os.environ["AZURE_OPENAI_API_KEY"])
    return _azure_client


def _image_block(img: Image.Image, max_edge: int) -> dict:
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": "image/jpeg",
            "data": _img_to_b64_jpeg(img, max_edge),
        },
    }


# ─── orientation detection ────────────────────────────────────────────────────

def detect_rotation(img: Image.Image) -> int:
    """Ask Haiku for the CW rotation (0/90/180/270) needed to upright the page.

    The checker's heuristic auto_rotate only handles portrait/landscape;
    it misses 180° flips that are common in scanned exam pages. Haiku is
    cheap enough to run per-page and reliable for simple orientation.
    """
    msg = client.messages.create(
        model=ORIENT_MODEL,
        max_tokens=10,
        messages=[{
            "role": "user",
            "content": [
                _image_block(img, ORIENT_MAX_EDGE),
                {"type": "text", "text": (
                    "זהו עמוד סרוק של מבחן מתמטיקה בכתב יד (עברית + נוסחאות). "
                    "בכמה מעלות עם כיוון השעון צריך לסובב אותו כדי שהטקסט יהיה זקוף וקריא? "
                    "ענה במספר אחד בלבד: 0, 90, 180 או 270."
                )},
            ],
        }],
    )
    txt = "".join(b.text for b in msg.content if b.type == "text")
    for cand in ("180", "270", "90", "0"):
        if cand in txt:
            return int(cand)
    return 0


def apply_rotation(img: Image.Image, deg: int) -> Image.Image:
    if deg == 0:
        return img
    return img.rotate(-deg, expand=True, resample=Image.BILINEAR, fillcolor=(255, 255, 255))


# ─── holistic math analysis ───────────────────────────────────────────────────

ANALYSIS_PROMPT = r"""אתה בודק מבחנים במתמטיקה וגאומטריה (שכבת חטיבה/תיכון) בישראל.
לפניך עמוד סרוק של פתרון תלמיד בכתב יד. נתח את העמוד כיחידה שלמה — אל תחתוך לשורות.
שים לב לסרטוטים, הוכחות גאומטריות ומעברים אלגבריים שפרושים על פני הדף.

כללים:
• כל ביטוי מתמטי ב-LaTeX תקני (תוכן בלבד, ללא $, למשל: x^2-6x+10 , \triangle ABC \cong \triangle DEF).
• אם צעד שגוי — ok=false והסבר ב-comment מה הטעות (חישובית / שיטתית / הגיונית / כתיב מתמטי).
• אם משהו לא קריא — כתוב "[לא קריא]" באותו שדה.
• verdict לפי הפתרון כולו: correct / partial / incorrect / unclear.
• feedback: משוב קצר ובונה לתלמיד בעברית (2-3 משפטים). כלול מה טוב ומה צריך שיפור.
• ניקוד מספרי:
  – points_max: הניקוד המקסימלי לתרגיל. אם סופקה רובריקה/מחוון — לפיה בדיוק.
    אחרת לפי הניקוד המודפס ליד התרגיל בדף, ואם אין — הערך סביר (ברירת מחדל 10).
  – points_earned: הניקוד שמגיע לתלמיד לפי איכות הפתרון. מספר בין 0 ל-points_max,
    חצאי נקודות מותרים. תן ניקוד חלקי הוגן לפתרון חלקי (אל תאפס בגלל טעות אחת).
• score_suggestion: נימוק קצר בעברית לניקוד שנתת, למשל "טעות חישוב בצעד האחרון (-3)"
  / "פתרון מלא ומנומק" / "לא ענה". זה נימוק, לא חובה לכלול בו מספר.
• id: מספר/אות התרגיל בדיוק כפי שמופיע בדף (למשל "1", "2א", "ב"). אם לא ברור — "?".
• topic: אחת מ: אלגברה / גאומטריה / חשבון / הסתברות / טריגונומטריה / מש"ב / אחר.

הניקוד שאתה מציע הוא **הצעה לבדיקת המורה** — המורה יאשר או יתקן. היה הוגן ועקבי."""

ANALYSIS_SCHEMA = {
    "type": "object",
    "properties": {
        "page_summary": {"type": "string"},
        "has_diagram": {"type": "boolean"},
        "diagram_description": {"type": "string"},
        "problems": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "id":               {"type": "string"},
                    "topic":            {"type": "string"},
                    "statement_latex":  {"type": "string"},
                    "student_steps": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "latex":   {"type": "string"},
                                "ok":      {"type": "boolean"},
                                "comment": {"type": "string"},
                            },
                            "required": ["latex", "ok", "comment"],
                            "additionalProperties": False,
                        },
                    },
                    "final_answer_latex": {"type": "string"},
                    "verdict":            {"type": "string"},
                    "points_max":         {"type": "number"},
                    "points_earned":      {"type": "number"},
                    "score_suggestion":   {"type": "string"},
                    "feedback":           {"type": "string"},
                },
                "required": [
                    "id", "topic", "statement_latex", "student_steps",
                    "final_answer_latex", "verdict", "points_max", "points_earned",
                    "score_suggestion", "feedback",
                ],
                "additionalProperties": False,
            },
        },
    },
    "required": ["page_summary", "has_diagram", "diagram_description", "problems"],
    "additionalProperties": False,
}


def _strip_additional_props(node):
    """Gemini's response_schema is an OpenAPI subset that rejects the
    additionalProperties key our Anthropic schema carries. Drop it recursively."""
    if isinstance(node, dict):
        return {k: _strip_additional_props(v) for k, v in node.items()
                if k != "additionalProperties"}
    if isinstance(node, list):
        return [_strip_additional_props(v) for v in node]
    return node


_GEMINI_SCHEMA = _strip_additional_props(ANALYSIS_SCHEMA)

# Groq / Azure are OpenAI-compatible and only support response_format=
# json_object (no schema), so the JSON shape is spelled out in the user turn.
_JSON_CONTRACT = (
    " החזר JSON בלבד, ללא טקסט נוסף, במבנה: "
    '{"page_summary":"...","has_diagram":true,"diagram_description":"...",'
    '"problems":[{"id":"1","topic":"אלגברה","statement_latex":"...",'
    '"student_steps":[{"latex":"...","ok":true,"comment":"..."}],'
    '"final_answer_latex":"...","verdict":"correct|partial|incorrect|unclear",'
    '"points_max":10,"points_earned":7,"score_suggestion":"...","feedback":"..."}]}'
)

_ANALYSIS_USER_TURN = "נתח את עמוד הפתרון וחזור JSON."


def _rubric_block(rubric: str) -> str:
    """Per-request rubric text, appended to the user turn (not the cached system
    prompt) so the teacher's grading key conditions scoring without busting the
    prompt cache."""
    rubric = (rubric or "").strip()
    if not rubric:
        return ""
    return (
        "\n\nרובריקה / מחוון מהמורה — נקד לפיה בדיוק (points_max ו-points_earned "
        "לכל תרגיל יתאימו לה):\n" + rubric
    )


def analyze_page(img: Image.Image, model_key: str = DEFAULT_MODEL,
                 rubric: str = "") -> dict:
    """Holistic full-page math analysis. Dispatches to the chosen provider;
    every backend returns the same structured JSON (ANALYSIS_SCHEMA shape).
    An optional rubric conditions the numeric scoring."""
    if model_key in _ANTHROPIC_IDS:
        return _analyze_anthropic(img, _ANTHROPIC_IDS[model_key], rubric)
    if model_key in _GEMINI_IDS:
        return _analyze_gemini(img, _GEMINI_IDS[model_key], rubric)
    if model_key in _GROQ_IDS:
        return _analyze_groq(img, _GROQ_IDS[model_key], rubric)
    if model_key in _AZURE_KEYS:
        return _analyze_azure(img, rubric)
    return _analyze_anthropic(img, _ANTHROPIC_IDS[DEFAULT_MODEL], rubric)


def _analyze_anthropic(img: Image.Image, model_id: str, rubric: str = "") -> dict:
    response = client.messages.create(
        model=model_id,
        max_tokens=6000,
        system=[{
            "type": "text",
            "text": ANALYSIS_PROMPT,
            "cache_control": {"type": "ephemeral"},
        }],
        output_config={
            "format": {"type": "json_schema", "schema": ANALYSIS_SCHEMA}
        },
        messages=[{
            "role": "user",
            "content": [
                _image_block(img, MAX_EDGE),
                {"type": "text", "text": _ANALYSIS_USER_TURN + _rubric_block(rubric)},
            ],
        }],
    )
    text = next(b.text for b in response.content if b.type == "text")
    return json.loads(text)


def _analyze_gemini(img: Image.Image, model_id: str, rubric: str = "") -> dict:
    from google.genai import types
    response = _gemini().models.generate_content(
        model=model_id,
        contents=[
            types.Part.from_bytes(data=_jpeg_bytes(img, MAX_EDGE), mime_type="image/jpeg"),
            _ANALYSIS_USER_TURN + _rubric_block(rubric),
        ],
        config=types.GenerateContentConfig(
            system_instruction=ANALYSIS_PROMPT,
            response_mime_type="application/json",
            response_schema=_GEMINI_SCHEMA,
            # Gemini 2.5 thinking tokens count against this budget; keep it
            # generous or the JSON gets truncated mid-string before it closes.
            max_output_tokens=24000,
        ),
    )
    return json.loads(response.text)


def _analyze_openai_compatible(client_obj, model_id: str, img: Image.Image,
                               rubric: str = "") -> dict:
    """Shared call shape for Groq and Azure (both OpenAI-compatible vision)."""
    b64 = base64.standard_b64encode(_jpeg_bytes(img, MAX_EDGE)).decode()
    response = client_obj.chat.completions.create(
        model=model_id,
        messages=[
            {"role": "system", "content": ANALYSIS_PROMPT},
            {"role": "user", "content": [
                {"type": "image_url",
                 "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                {"type": "text",
                 "text": "נתח את עמוד הפתרון." + _rubric_block(rubric) + _JSON_CONTRACT},
            ]},
        ],
        response_format={"type": "json_object"},
        max_tokens=8000,
        temperature=0,
    )
    return json.loads(response.choices[0].message.content)


def _analyze_groq(img: Image.Image, model_id: str, rubric: str = "") -> dict:
    return _analyze_openai_compatible(_groq(), model_id, img, rubric)


def _analyze_azure(img: Image.Image, rubric: str = "") -> dict:
    return _analyze_openai_compatible(
        _azure(), os.environ["AZURE_OPENAI_DEPLOYMENT"], img, rubric)


# ─── error handling ───────────────────────────────────────────────────────────

def humanize_error(exc: Exception) -> dict:
    """Turn an SDK exception into a user-readable Hebrew error payload."""
    raw = (str(exc) or repr(exc)).strip()
    tech = raw[:2000]

    if isinstance(exc, anthropic.RateLimitError):
        return {
            "message": "חרגת ממכסת Claude לרגע זה",
            "details": "Anthropic מגביל בקשות לדקה (RPM) וטוקנים לדקה (TPM). נסה שוב בעוד דקה.",
            "technical": tech,
        }
    if isinstance(exc, anthropic.AuthenticationError):
        return {
            "message": "מפתח Anthropic API לא תקף",
            "details": "ודא שהמשתנה ANTHROPIC_API_KEY מוגדר נכון ב-.env ושהמפתח עדיין פעיל.",
            "technical": tech,
        }
    if isinstance(exc, (anthropic.APITimeoutError, anthropic.APIConnectionError)):
        return {
            "message": "החיבור ל-Anthropic נכשל",
            "details": "ייתכן שיש בעיית רשת או ש-Anthropic לא זמין כרגע. נסה שוב בעוד דקה.",
            "technical": tech,
        }
    if isinstance(exc, KeyError) and exc.args and str(exc.args[0]) == "ANTHROPIC_API_KEY":
        return {
            "message": "מפתח Anthropic API חסר",
            "details": "המשתנה ANTHROPIC_API_KEY לא מוגדר. הוסף אותו ל-.env ואז restart.",
            "technical": tech,
        }
    if isinstance(exc, (json.JSONDecodeError, ValueError)):
        return {
            "message": "המודל החזיר תשובה לא תקפה",
            "details": "לא הצלחנו לפענח את תשובת המודל. זו בדרך כלל בעיה זמנית — נסה שוב.",
            "technical": tech,
        }
    if isinstance(exc, anthropic.APIStatusError):
        return {
            "message": f"שגיאת API ({getattr(exc, 'status_code', '?')})",
            "details": "השירות החזיר תשובה לא תקינה. נסה שוב.",
            "technical": tech,
        }
    return {
        "message": "תקלה לא צפויה",
        "details": "אירעה שגיאה לא מזוהה. נסה שוב; אם זה חוזר צרף את הפרטים הטכניים.",
        "technical": tech,
    }


# ─── streaming job runner ─────────────────────────────────────────────────────

STAGE_LABELS = {
    "orient":  "מזהה סיבוב (Haiku)",
    "analyze": "מנתח מתמטיקה",
}
_STEPS_PER_PAGE = 2  # orient, analyze


def _evict_stale_jobs():
    now = time.time()
    for jid in [k for k, v in JOBS.items() if now - v.get("ts", 0) > JOB_TTL]:
        JOBS.pop(jid, None)


def _process_stream(file_bytes: bytes, ext: str, auto_orient: bool,
                    model_key: str = DEFAULT_MODEL, model_label: str = "",
                    rubric: str = ""):
    """Generator: yields SSE-ready progress dicts, then a final result dict."""
    pages_imgs = file_to_pages(file_bytes, ext)
    total = len(pages_imgs)
    short_label = (model_label.split("·")[0].strip() or "מודל")
    log.info("[job] rendered %d page(s) at %d DPI; model=%s auto_orient=%s",
             total, RENDER_DPI, model_key, auto_orient)

    def progress(stage: str, page: int, completed: int) -> dict:
        # Emitted *before* each stage runs, so the label reflects the stage
        # currently executing (the analysis pass is the slow one).
        label = STAGE_LABELS.get(stage, stage)
        if stage == "analyze":
            label = f"{label} ({short_label})"
        return {
            "type": "progress",
            "page": page,
            "total_pages": total,
            "stage": stage,
            "label": label,
            "pct": round(100 * completed / max(1, total * _STEPS_PER_PAGE)),
        }

    results = []
    imgs = []  # oriented analysis-resolution images, kept server-side for re-analysis
    for idx, img in enumerate(pages_imgs):
        p = idx + 1
        base = idx * _STEPS_PER_PAGE

        yield progress("orient", p, base)
        if auto_orient:
            t0 = time.time()
            rotation = detect_rotation(img)
            img = apply_rotation(img, rotation)
            log.info("[page %d] orient → %d° (%.1fs)", p, rotation, time.time() - t0)
        else:
            rotation = 0

        yield progress("analyze", p, base + 1)
        img = _downscale(img, MAX_EDGE)
        t0 = time.time()
        analysis = analyze_page(img, model_key, rubric)
        log.info(
            "[page %d] model=%s rotation=%d° problems=%d verdicts=%s (%.1fs)",
            p, model_key, rotation,
            len(analysis.get("problems", [])),
            [pr.get("verdict") for pr in analysis.get("problems", [])],
            time.time() - t0,
        )

        results.append({
            "page": p,
            "rotation_applied": rotation,
            "image_b64": _img_to_b64_jpeg(img, PREVIEW_MAX_EDGE, quality=85),
            "analysis": analysis,
        })
        imgs.append(img)

    yield {"type": "result", "pages": results, "imgs": imgs}


def _run_job(job_id: str, file_bytes: bytes, ext: str, auto_orient: bool, filename: str,
             model_key: str = DEFAULT_MODEL, model_label: str = "", rubric: str = ""):
    job = JOBS[job_id]
    q: queue.Queue = job["q"]
    try:
        for ev in _process_stream(file_bytes, ext, auto_orient, model_key,
                                  model_label, rubric):
            if ev["type"] == "result":
                job["pages"] = ev["pages"]
                job["imgs"] = ev["imgs"]
                job["filename"] = filename
                q.put({"type": "done"})
            else:
                q.put(ev)
    except Exception as e:
        log.exception("_run_job failed for job %s", job_id)
        q.put({"type": "error", **humanize_error(e)})
    finally:
        q.put({"type": "end"})


# ─── routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    # no-store: the UI markup changes often during development and a stale
    # cached index.html silently sends old form values (e.g. auto_orient
    # defaulting on). Force the browser to always fetch fresh HTML.
    resp = Response(render_template("index.html"))
    resp.headers["Cache-Control"] = "no-store, max-age=0"
    return resp


@app.route("/analyze/start", methods=["POST"])
def analyze_start():
    _evict_stale_jobs()
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "לא הועלה קובץ"}), 400
    ext = os.path.splitext(f.filename or "")[1].lower()
    if ext not in ACCEPTED_EXTS:
        return jsonify({"error": f"סוג קובץ לא נתמך: {ext}"}), 400

    file_bytes = f.read()
    auto_orient = request.form.get("auto_orient", "0") == "1"
    model_key, model_label = resolve_model(request.form.get("model"))
    rubric = (request.form.get("rubric") or "").strip()
    filename = f.filename or "תרגיל"

    job_id = str(uuid.uuid4())
    log.info(
        "[upload] job=%s file=%r ext=%s size=%.1fKB model=%s auto_orient=%s rubric=%dch",
        job_id[:8], filename, ext, len(file_bytes) / 1024, model_key, auto_orient,
        len(rubric),
    )
    JOBS[job_id] = {"q": queue.Queue(), "ts": time.time(), "pages": None,
                    "imgs": None, "filename": None, "model": model_key,
                    "rubric": rubric}
    threading.Thread(
        target=_run_job,
        args=(job_id, file_bytes, ext, auto_orient, filename, model_key,
              model_label, rubric),
        daemon=True,
    ).start()
    return jsonify({"job_id": job_id})


@app.route("/analyze/stream/<job_id>")
def analyze_stream(job_id: str):
    job = JOBS.get(job_id)
    if not job:
        return jsonify({"error": "job not found"}), 404

    def gen():
        q: queue.Queue = job["q"]
        while True:
            ev = q.get()
            yield f"data: {json.dumps(ev, ensure_ascii=False)}\n\n"
            if ev["type"] in ("end", "error"):
                break

    return Response(
        gen(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/reanalyze", methods=["POST"])
def reanalyze():
    """Re-run analysis on a single page after the user manually rotated it.

    Body: {job_id, page, rotate}. `rotate` is extra CW degrees (0/90/180/270)
    to apply to the page's current orientation before re-analyzing.
    """
    data = request.get_json(silent=True) or {}
    job = JOBS.get(data.get("job_id"))
    if not job or not job.get("imgs"):
        return jsonify({"message": "התרגיל לא נמצא", "details": "ייתכן שהזמן הקצוב פג. טען מחדש."}), 404

    try:
        idx = int(data.get("page", 0)) - 1
    except (TypeError, ValueError):
        idx = -1
    if not (0 <= idx < len(job["imgs"])):
        return jsonify({"message": "מספר עמוד לא תקין"}), 400

    rotate = int(data.get("rotate", 0)) % 360
    img = apply_rotation(job["imgs"][idx], rotate) if rotate else job["imgs"][idx]

    model_key = job.get("model", DEFAULT_MODEL)
    try:
        analysis = analyze_page(img, model_key, job.get("rubric", ""))
    except Exception as e:
        log.exception("reanalyze failed")
        return jsonify(humanize_error(e)), 502

    new_rot = (job["pages"][idx].get("rotation_applied", 0) + rotate) % 360
    page_data = {
        "page": idx + 1,
        "rotation_applied": new_rot,
        "image_b64": _img_to_b64_jpeg(img, PREVIEW_MAX_EDGE, quality=85),
        "analysis": analysis,
    }
    job["imgs"][idx] = img
    job["pages"][idx] = page_data
    job["ts"] = time.time()
    log.info("[reanalyze page %d] rotate=%d° → rotation=%d° problems=%d",
             idx + 1, rotate, new_rot, len(analysis.get("problems", [])))
    return jsonify(page_data)


# fields the teacher may edit; everything else (image, latex, steps) stays as
# the model produced it and is never overwritten by a client update.
_EDITABLE_FIELDS = ("points_earned", "points_max", "verdict",
                    "score_suggestion", "feedback")


@app.route("/update/<job_id>", methods=["POST"])
def update_result(job_id: str):
    """Persist the teacher's edits (per-problem scores/feedback + approval) back
    onto the job so the exports reflect what the teacher confirmed, not the raw
    AI suggestion. Merges only whitelisted fields by position; the scan image
    and the model's transcription/steps are left intact."""
    job = JOBS.get(job_id)
    if not job or not job.get("pages"):
        return jsonify({"error": "תוצאה לא נמצאה"}), 404

    data = request.get_json(silent=True) or {}
    incoming = data.get("pages") or []
    for pi, page in enumerate(job["pages"]):
        if pi >= len(incoming):
            break
        in_problems = ((incoming[pi] or {}).get("analysis") or {}).get("problems") or []
        problems = (page.get("analysis") or {}).get("problems") or []
        for qi, pr in enumerate(problems):
            if qi >= len(in_problems):
                break
            src = in_problems[qi] or {}
            for fld in _EDITABLE_FIELDS:
                if fld in src:
                    pr[fld] = src[fld]

    if "approved" in data:
        job["approved"] = bool(data["approved"])
    job["ts"] = time.time()
    log.info("[update] job=%s approved=%s", job_id[:8], job.get("approved"))
    return jsonify({"ok": True, "approved": job.get("approved", False)})


@app.route("/result-data/<job_id>")
def result_data(job_id: str):
    """Return the analysis result as JSON (for client-side rendering after SSE)."""
    job = JOBS.get(job_id)
    if not job or not job.get("pages"):
        return jsonify({"error": "תוצאה לא נמצאה"}), 404
    return jsonify({"pages": job["pages"], "filename": job.get("filename", ""),
                    "approved": job.get("approved", False)})


# ─── export builders (HTML + Word) ──────────────────────────────────────────

VERDICT_HE = {"correct": "נכון ✓", "partial": "חלקי", "incorrect": "שגוי ✗", "unclear": "לא ברור"}
VERDICT_COLOR = {"correct": "1e7e34", "partial": "a0740a", "incorrect": "b54343", "unclear": "5a6b82"}


def _num(v):
    """Coerce a points value to a number, or None if missing/blank/non-numeric."""
    if v is None or v == "":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def _fmt_pts(v) -> str:
    """Render a points number without a trailing .0 (7.0 → "7", 7.5 → "7.5")."""
    n = _num(v)
    if n is None:
        return "—"
    return str(int(n)) if n == int(n) else str(n)


def compute_totals(pages: list[dict]) -> tuple[float, float]:
    """Sum (earned, max) points across every problem on every page."""
    earned = total = 0.0
    for p in pages or []:
        for pr in ((p.get("analysis") or {}).get("problems") or []):
            mx = _num(pr.get("points_max"))
            if mx is None:
                continue
            total += mx
            earned += _num(pr.get("points_earned")) or 0
    return earned, total


def build_result_html(pages: list[dict], filename: str, approved: bool = False) -> str:
    """Standalone, self-contained HTML report. KaTeX is pulled from a CDN and
    auto-renders \\(...\\) so the LaTeX shows as real math; the scan images are
    embedded as base64 so the file opens anywhere with no server. Mirrors the
    web UI layout (verdict badges, ok/bad step coloring, comments, feedback)."""
    import html as _html

    def esc(s) -> str:
        return _html.escape(str(s if s is not None else ""))

    def math(latex) -> str:
        # KaTeX auto-render reads textContent, so escaping is safe (and needed
        # for inequalities like a<b). Wrapped in \( \) inline delimiters.
        return r"\(" + esc(latex) + r"\)"

    body: list[str] = []
    for p in pages or []:
        a = p.get("analysis") or {}
        rot = p.get("rotation_applied") or 0
        rot_lbl = f" · סובב {rot}°" if rot else ""
        body.append(f'<section class="page"><h2>עמוד {esc(p.get("page", "?"))}'
                    f'<span class="rot">{esc(rot_lbl)}</span></h2><div class="layout">')
        if p.get("image_b64"):
            body.append(f'<div class="imgwrap"><img src="data:image/jpeg;base64,'
                        f'{p["image_b64"]}" alt="עמוד {esc(p.get("page", ""))}"></div>')
        body.append('<div class="analysis">')
        if a.get("page_summary"):
            body.append(f'<div class="summary">{esc(a["page_summary"])}</div>')
        if a.get("has_diagram"):
            body.append(f'<div class="diagram">▣ סרטוט: {esc(a.get("diagram_description", ""))}</div>')
        problems = a.get("problems") or []
        if not problems:
            body.append('<div class="noprob">לא זוהו תרגילים בעמוד זה.</div>')
        for pr in problems:
            v = pr.get("verdict", "unclear")
            body.append('<div class="prob"><div class="prob-head">'
                        f'<span class="prob-id">תרגיל {esc(pr.get("id", "?"))}</span>'
                        f'<span class="topic">{esc(pr.get("topic", ""))}</span>'
                        f'<span class="badge b-{esc(v)}">{esc(VERDICT_HE.get(v, v))}</span></div>')
            if pr.get("statement_latex"):
                body.append(f'<div class="step statement" dir="ltr">{math(pr["statement_latex"])}</div>')
            for s in pr.get("student_steps") or []:
                cls = "bad" if s.get("ok") is False else ("ok" if s.get("ok") is True else "")
                body.append(f'<div class="step {cls}" dir="ltr">{math(s.get("latex", ""))}')
                if s.get("comment"):
                    body.append(f'<div class="step-comment">⚠ {esc(s["comment"])}</div>')
                body.append('</div>')
            if pr.get("final_answer_latex"):
                body.append(f'<div class="step final" dir="ltr"><b>תשובה:</b> {math(pr["final_answer_latex"])}</div>')
            if _num(pr.get("points_max")) is not None:
                body.append(
                    '<div class="score-row"><span class="pts">ניקוד: '
                    f'{_fmt_pts(pr.get("points_earned"))} / {_fmt_pts(pr.get("points_max"))}'
                    '</span>'
                    + (f'<span class="score-note">{esc(pr["score_suggestion"])}</span>'
                       if pr.get("score_suggestion") else "")
                    + '</div>')
            elif pr.get("score_suggestion"):
                body.append(f'<div class="score-row">{esc(pr["score_suggestion"])}</div>')
            if pr.get("feedback"):
                body.append(f'<div class="feedback">{esc(pr["feedback"])}</div>')
            body.append('</div>')
        body.append('</div></div></section>')

    head = """<!DOCTYPE html>
<html dir="rtl" lang="he"><head><meta charset="utf-8">
<title>בדיקת מתמטיקה — __TITLE__</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"
  onload="renderMathInElement(document.body,{delimiters:[{left:'\\\\(',right:'\\\\)',display:false},{left:'\\\\[',right:'\\\\]',display:true}],throwOnError:false});"></script>
<style>
  body{font-family:'Segoe UI',Arial,sans-serif;color:#2b3442;background:#f4f1ea;margin:0;padding:2rem;line-height:1.5}
  .doc-title{font-size:1.5rem;font-weight:800;color:#2e7286;margin:0 0 0.3rem}
  .doc-meta{color:#6e6a5c;margin:0 0 1.5rem;font-size:0.9rem}
  .page{background:#fff;border:1px solid #d8cfb6;border-radius:12px;padding:1.2rem 1.4rem;margin-bottom:1.5rem;box-shadow:0 1px 4px rgba(0,0,0,.06)}
  .page h2{margin:0 0 1rem;color:#2e7286;font-size:1.2rem}
  .rot{color:#6e6a5c;font-size:0.85rem;font-weight:400}
  .layout{display:flex;gap:1.4rem;align-items:flex-start;flex-wrap:wrap}
  .imgwrap{flex:0 0 320px;max-width:340px}
  .imgwrap img{width:100%;border:1px solid #d8cfb6;border-radius:8px}
  .analysis{flex:1;min-width:300px}
  .summary{font-size:0.95rem;color:#3a4250;margin-bottom:0.6rem}
  .diagram{font-size:0.85rem;color:#2e7286;background:#eef4f7;border-radius:6px;padding:0.3rem 0.6rem;margin-bottom:0.6rem;display:inline-block}
  .noprob{color:#6e6a5c;font-style:italic}
  .prob{border:1px solid #e7e0cd;border-radius:10px;padding:0.8rem 1rem;margin-bottom:0.9rem}
  .prob-head{display:flex;gap:0.6rem;align-items:center;margin-bottom:0.5rem;flex-wrap:wrap}
  .prob-id{font-weight:700}
  .topic{font-size:0.78rem;color:#6e6a5c;background:#ece8db;border-radius:999px;padding:0.15rem 0.6rem}
  .badge{font-size:0.82rem;padding:0.18rem 0.7rem;border-radius:999px;font-weight:700}
  .b-correct{background:rgba(30,126,52,.15);color:#1e7e34}
  .b-partial{background:rgba(160,116,10,.15);color:#a0740a}
  .b-incorrect{background:rgba(181,67,67,.15);color:#b54343}
  .b-unclear{background:rgba(90,107,130,.15);color:#5a6b82}
  .step{padding:0.5rem 0.8rem;margin:0.3rem 0;border-radius:6px;border-inline-start:3px solid #d8cfb6;background:#f6f3eb;direction:ltr;text-align:left}
  .step.ok{border-inline-start-color:#1e7e34}
  .step.bad{border-inline-start-color:#b54343;background:#fff0f0}
  .step.statement{border-inline-start-color:#2e7286;background:#eef4f7}
  .step.final{border-inline-start-color:#b3924a;background:#faf5e9;font-weight:600}
  .step-comment{color:#b54343;font-size:0.85rem;margin-top:0.3rem;direction:rtl;text-align:right;font-weight:500}
  .score-row{color:#3a4250;font-size:0.9rem;margin:0.6rem 0 0.3rem;display:flex;gap:0.6rem;align-items:baseline;flex-wrap:wrap}
  .score-row .pts{font-weight:700;color:#2e7286}
  .score-row .score-note{color:#6e6a5c;font-size:0.85rem}
  .feedback{margin-top:0.5rem;padding-top:0.6rem;border-top:1px dashed #d8cfb6;font-size:0.93rem}
  .total-box{background:#fff;border:1px solid #d8cfb6;border-radius:12px;padding:1rem 1.4rem;margin-bottom:1.5rem;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:0.8rem}
  .total-grade{font-size:1.3rem;font-weight:800;color:#2e7286}
  .approve-badge{font-size:0.9rem;font-weight:700;padding:0.3rem 0.9rem;border-radius:999px}
  .approve-yes{background:rgba(30,126,52,.15);color:#1e7e34}
  .approve-no{background:rgba(160,116,10,.15);color:#a0740a}
  @media print{body{background:#fff;padding:0}.page{box-shadow:none;break-inside:avoid}}
</style></head><body>
<div class="doc-title">בדיקת מתמטיקה — השכלה</div>
<div class="doc-meta">קובץ: __TITLE__ · __NPAGES__ עמודים</div>
__TOTAL_BOX__
"""
    npages = len(pages or [])
    earned, total = compute_totals(pages)
    if total > 0:
        badge = ('<span class="approve-badge approve-yes">✓ אושר ע"י המורה</span>'
                 if approved else
                 '<span class="approve-badge approve-no">טיוטה — טרם אושר</span>')
        total_box = (
            '<div class="total-box">'
            f'<span class="total-grade">ציון כולל: {_fmt_pts(earned)} / {_fmt_pts(total)}</span>'
            f'{badge}</div>')
    else:
        total_box = ""
    head = (head.replace("__TITLE__", _html.escape(filename))
                .replace("__NPAGES__", str(npages))
                .replace("__TOTAL_BOX__", total_box))
    return head + "\n".join(body) + "\n</body></html>"


def build_result_docx(pages: list[dict], filename: str, approved: bool = False) -> bytes:
    """Word (.docx) report. python-docx can't typeset LaTeX, so math is kept as
    monospace LTR text (the strings are simple, e.g. x^2-6x+10) — honest and
    editable. Each page carries its scan image, verdicts, comments, score and
    feedback. Hebrew paragraphs are right-aligned + bidi for correct RTL flow."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def rtl(par):
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = par._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
        return par

    def mono(run):
        run.font.name = "Consolas"
        run.font.size = Pt(10.5)
        return run

    doc = Document()
    rtl(doc.add_heading("בדיקת מתמטיקה — השכלה", level=1))
    meta = rtl(doc.add_paragraph())
    meta.add_run(f"קובץ: {filename}").bold = True

    earned, total = compute_totals(pages)
    if total > 0:
        tp = rtl(doc.add_paragraph())
        tr = tp.add_run(f"ציון כולל: {_fmt_pts(earned)} / {_fmt_pts(total)}")
        tr.bold = True
        tr.font.size = Pt(14)
        tr.font.color.rgb = RGBColor(0x2E, 0x72, 0x86)
        sp = rtl(doc.add_paragraph())
        sr = sp.add_run('✓ אושר ע"י המורה' if approved else "טיוטה — טרם אושר ע\"י המורה")
        sr.bold = True
        sr.font.color.rgb = RGBColor(0x1E, 0x7E, 0x34) if approved else RGBColor(0xA0, 0x74, 0x0A)

    for idx, p in enumerate(pages or []):
        if idx > 0:
            doc.add_page_break()
        a = p.get("analysis") or {}
        rot = p.get("rotation_applied") or 0
        rot_lbl = f"  (סובב {rot}°)" if rot else ""
        rtl(doc.add_heading(f"עמוד {p.get('page', idx + 1)}{rot_lbl}", level=2))

        if p.get("image_b64"):
            try:
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(
                    io.BytesIO(base64.b64decode(p["image_b64"])), width=Inches(3.2))
            except Exception:
                log.warning("docx: failed to embed image for page %s", p.get("page"))

        if a.get("page_summary"):
            rtl(doc.add_paragraph(a["page_summary"]))
        if a.get("has_diagram"):
            d = rtl(doc.add_paragraph())
            r = d.add_run(f"▣ סרטוט: {a.get('diagram_description', '')}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x2E, 0x72, 0x86)

        problems = a.get("problems") or []
        if not problems:
            ip = rtl(doc.add_paragraph())
            ip.add_run("לא זוהו תרגילים בעמוד זה.").italic = True

        for pr in problems:
            v = pr.get("verdict", "unclear")
            h = rtl(doc.add_heading(level=3))
            hr = h.add_run(f"תרגיל {pr.get('id', '?')}  ·  {pr.get('topic', '')}  ·  "
                           f"{VERDICT_HE.get(v, v)}")
            hr.font.color.rgb = RGBColor.from_string(VERDICT_COLOR.get(v, "5a6b82"))

            if pr.get("statement_latex"):
                sp = doc.add_paragraph()
                sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sp.add_run("נתון: ").bold = True
                mono(sp.add_run(pr["statement_latex"]))

            for s in pr.get("student_steps") or []:
                mark = "✗ " if s.get("ok") is False else ("✓ " if s.get("ok") is True else "• ")
                stp = doc.add_paragraph()
                stp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                mk = stp.add_run(mark)
                mk.bold = True
                mk.font.color.rgb = RGBColor(0xB5, 0x43, 0x43) if s.get("ok") is False \
                    else (RGBColor(0x1E, 0x7E, 0x34) if s.get("ok") is True else RGBColor(0x5A, 0x6B, 0x82))
                mono(stp.add_run(s.get("latex", "")))
                if s.get("comment"):
                    cp = rtl(doc.add_paragraph())
                    cr = cp.add_run(f"⚠ {s['comment']}")
                    cr.italic = True
                    cr.font.size = Pt(9.5)
                    cr.font.color.rgb = RGBColor(0xB5, 0x43, 0x43)

            if pr.get("final_answer_latex"):
                fp = doc.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fp.add_run("תשובה: ").bold = True
                mono(fp.add_run(pr["final_answer_latex"]))

            if _num(pr.get("points_max")) is not None:
                scp = rtl(doc.add_paragraph())
                scp.add_run(
                    f"ניקוד: {_fmt_pts(pr.get('points_earned'))} / "
                    f"{_fmt_pts(pr.get('points_max'))}").bold = True
                if pr.get("score_suggestion"):
                    note = scp.add_run(f"   ({pr['score_suggestion']})")
                    note.italic = True
                    note.font.size = Pt(9.5)
                    note.font.color.rgb = RGBColor(0x6E, 0x6A, 0x5C)
            elif pr.get("score_suggestion"):
                rtl(doc.add_paragraph()).add_run(pr["score_suggestion"]).bold = True
            if pr.get("feedback"):
                fb = rtl(doc.add_paragraph())
                fb.add_run(pr["feedback"]).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/save/<job_id>")
def save_json(job_id: str):
    """Download the analysis result as a JSON file."""
    job = JOBS.get(job_id)
    if not job or not job.get("pages"):
        return "תוצאה לא נמצאה", 404
    filename = job.get("filename", "תרגיל")
    base = os.path.splitext(filename)[0]
    today = datetime.date.today().strftime("%Y%m%d")
    save_name = f"{base}_math_{today}.json"
    payload = json.dumps(
        {"pages": job["pages"], "filename": filename},
        ensure_ascii=False,
        indent=2,
    )
    return Response(
        payload,
        mimetype="application/json",
        headers={"Content-Disposition": f'attachment; filename="{save_name}"'},
    )


def _export_basename(job: dict) -> str:
    base = os.path.splitext(job.get("filename", "תרגיל") or "תרגיל")[0]
    today = datetime.date.today().strftime("%Y%m%d")
    return f"{base}_math_{today}"


@app.route("/save/html/<job_id>")
def save_html(job_id: str):
    """Download the analysis result as a standalone HTML report (KaTeX math)."""
    job = JOBS.get(job_id)
    if not job or not job.get("pages"):
        return "תוצאה לא נמצאה", 404
    html_doc = build_result_html(job["pages"], job.get("filename", "תרגיל"),
                                 job.get("approved", False))
    return Response(
        html_doc,
        mimetype="text/html",
        headers={"Content-Disposition":
                 f'attachment; filename="{_export_basename(job)}.html"'},
    )


@app.route("/save/docx/<job_id>")
def save_docx(job_id: str):
    """Download the analysis result as a Word (.docx) document."""
    job = JOBS.get(job_id)
    if not job or not job.get("pages"):
        return "תוצאה לא נמצאה", 404
    docx_bytes = build_result_docx(job["pages"], job.get("filename", "תרגיל"),
                                   job.get("approved", False))
    return Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":
                 f'attachment; filename="{_export_basename(job)}.docx"'},
    )


@app.route("/load", methods=["POST"])
def load_json():
    """Parse and return a previously saved JSON result file."""
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "לא הועלה קובץ"}), 400
    try:
        data = json.loads(f.read().decode("utf-8"))
        return jsonify({
            "pages": data.get("pages", []),
            "filename": data.get("filename", "שמור.json"),
        })
    except (json.JSONDecodeError, ValueError) as e:
        return jsonify({"error": f"קובץ JSON לא תקין: {e}"}), 400


if __name__ == "__main__":
    app.run(debug=True, port=5051, threaded=True)
