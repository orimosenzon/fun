import base64
import difflib
import hashlib
import hmac
import io
import json
import logging
import os
import queue
import re
import threading
import uuid

import anthropic
import fitz
import numpy as np
from dotenv import load_dotenv
from flask import (
    Flask,
    Response,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
)
from PIL import Image, ImageOps

# override=True: .env is the local source of truth for keys. Without it a
# stale GEMINI_API_KEY/ANTHROPIC_API_KEY left in the shell shadows the real
# one. On HF there's no .env, so this is a no-op and the Secret wins.
load_dotenv(override=True)

# Pipeline diagnostics go to a file (not just stdout) so the segmentation
# behaviour can be inspected after a run without watching the console.
LOG_PATH = os.path.join(os.path.dirname(__file__), "haskala.log")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(message)s",
    handlers=[logging.FileHandler(LOG_PATH, mode="a"), logging.StreamHandler()],
)
log = logging.getLogger("haskala")

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024

# Set by desktop.py. When true we're a local single-user native app, so the
# file-path endpoints (reading an arbitrary path off disk) are safe. When
# false we're a public web server and those endpoints are a file-read hole,
# so they're refused — the browser-upload paths cover the web case.
DESKTOP_MODE = os.environ.get("HASKALA_DESKTOP") == "1"

# Public deployments must set HASKALA_USER / HASKALA_PASS so OCR (which
# spends Anthropic credit) sits behind Basic Auth. Locally / on desktop
# they're unset and auth is skipped.
BASIC_USER = os.environ.get("HASKALA_USER")
BASIC_PASS = os.environ.get("HASKALA_PASS")
if not DESKTOP_MODE and not (BASIC_USER and BASIC_PASS):
    logging.getLogger("haskala").warning(
        "running without Basic Auth — set HASKALA_USER/HASKALA_PASS "
        "before exposing this publicly"
    )


@app.before_request
def _require_basic_auth():
    if DESKTOP_MODE or not (BASIC_USER and BASIC_PASS):
        return  # local desktop or dev: no gate
    if request.path.startswith("/static/"):
        return
    auth = request.authorization
    if (
        auth
        and auth.type == "basic"
        and hmac.compare_digest(auth.username or "", BASIC_USER)
        and hmac.compare_digest(auth.password or "", BASIC_PASS)
    ):
        return
    return Response(
        "נדרשת הזדהות",
        401,
        {"WWW-Authenticate": 'Basic realm="haskala"'},
    )


client = anthropic.Anthropic()

# Render PDF at this DPI. Higher = better OCR but more tokens/cost.
# 200 DPI gives a good balance — A4 → ~1700×2340 px, well under Opus 4.7's 2576px limit.
RENDER_DPI = 200

OCR_SYSTEM_PROMPT = """אתה מבצע OCR מדויק על צילום של תרגיל כתוב ביד.

הקלט עשוי להיות בעברית, בערבית, באנגלית, או שילוב של כמה שפות באותו דף
(למשל הוראות בעברית ותשובות באנגלית). זהה את הכתב של כל שורה בנפרד ותמלל
אותו כפי שהוא — אל תתרגם ואל תמיר בין סקריפטים.

כללי תמלול קריטיים:
1. תמלל בדיוק את מה שכתוב — אל תתקן שום דבר.
2. שמור על שגיאות כתיב בדיוק כפי שהן.
3. כל שורה ויזואלית בתמונה = פריט אחד בפלט, מלמעלה למטה לפי הסדר.
4. שמור על כל סימני הפיסוק בדיוק כפי שנכתבו (כולל פיסוק חסר).
5. שמור על אותיות גדולות וקטנות באנגלית בדיוק כפי שנכתבו.
6. אם תו לא ברור, תמלל את מה שהכי דומה לצורה הכתובה.
7. אם מילה לא קריאה לחלוטין, כתוב במקומה: [לא קריא]
8. שורות ריקות בתמונה — דלג עליהן (אל תפיק פריט לשורה ריקה).

החזר את התוצאה כ-JSON תקין במבנה: {"lines": [{"text": "..."}, ...]}."""

OCR_SCHEMA = {
    "type": "object",
    "properties": {
        "lines": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "text": {"type": "string"},
                },
                "required": ["text"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["lines"],
    "additionalProperties": False,
}


IMAGE_EXTS = (".jpg", ".jpeg", ".png")
ACCEPTED_EXTS = (".pdf",) + IMAGE_EXTS


def pdf_to_page_images(pdf_bytes: bytes) -> list[Image.Image]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    for page in doc:
        pix = page.get_pixmap(dpi=RENDER_DPI)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        images.append(img)
    doc.close()
    return images


def file_to_page_images(data: bytes, ext: str) -> list[Image.Image]:
    """A PDF → one rendered image per page; a photo (jpg/png) → one page.

    The rest of the pipeline (deskew → ocr → segment → crop) is identical
    for both — a phone photo is just a single-page document with no render
    step to do, so it skips straight to the same handling a PDF page gets.
    """
    if ext == ".pdf":
        return pdf_to_page_images(data)
    # Phone cameras store portrait shots as landscape pixels + an EXIF
    # Orientation tag; PIL ignores the tag by default, so the page comes
    # out rotated 90° unless we transpose here.
    img = ImageOps.exif_transpose(Image.open(io.BytesIO(data)))
    return [img.convert("RGB")]


def image_to_b64(img: Image.Image, fmt: str = "PNG") -> str:
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return base64.standard_b64encode(buf.getvalue()).decode("utf-8")


# Longest edge of the image actually sent to the model. A 200-DPI A4 PDF
# page is ~2339 px tall, which this pipeline is already tuned for, so a
# phone photo is brought down to the same ballpark. It also keeps us well
# under the API's 5 MB-per-image cap — a full-res photo re-encoded as PNG
# blows past that; downscaling + JPEG keeps it small with no OCR loss.
OCR_MAX_EDGE = 2400


def ocr_jpeg_bytes(img: Image.Image) -> bytes:
    """Downscaled JPEG bytes — what actually goes to whichever model."""
    long_edge = max(img.size)
    if long_edge > OCR_MAX_EDGE:
        scale = OCR_MAX_EDGE / long_edge
        img = img.resize(
            (round(img.width * scale), round(img.height * scale)),
            Image.LANCZOS,
        )
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    return buf.getvalue()


# A full-page reference image kept beside the cut strips. It's only a visual
# aid for the reviewer — when a line cut looks wrong she can glance at the
# whole page to see where the strip came from — not OCR input, so a smaller
# JPEG is plenty and keeps the saved JSON from ballooning. This is the *raw*
# uploaded page, before deskew: what the reviewer actually photographed.
ORIGINAL_MAX_EDGE = 1800


def original_preview_b64(img: Image.Image) -> str:
    """Downscaled JPEG of the raw page — a reviewing aid, not OCR input."""
    long_edge = max(img.size)
    if long_edge > ORIGINAL_MAX_EDGE:
        scale = ORIGINAL_MAX_EDGE / long_edge
        img = img.resize(
            (round(img.width * scale), round(img.height * scale)),
            Image.LANCZOS,
        )
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=85)
    return base64.standard_b64encode(buf.getvalue()).decode("utf-8")


_USER_TURN = "תמלל שורה-שורה והחזר JSON."

# Model picker. Keys are what the UI / request send; values are the human
# label shown in the dropdown. Adding a provider = one entry + one _ocr_*.
MODELS = {
    "claude": "Claude (Opus 4.7)",
    "gemini": "Gemini (2.5 Flash)",
    "gemini-lite": "Gemini (2.5 Flash-Lite)",
    "groq-scout": "Groq (Llama 4 Scout — חינמי)",
}
DEFAULT_MODEL = "claude"


@app.context_processor
def _inject_models():
    """Makes the model list available to every index.html render without
    threading it through each render_template call site."""
    return {"models": MODELS, "default_model": DEFAULT_MODEL}


@app.context_processor
def _inject_rubrics():
    """Initial rubric list for empty-state renders. Result renders pass the
    list via server_data instead so the UI can refresh after adding one."""
    return {
        "initial_rubrics": list_rubrics(),
        "default_rubric_id": DEFAULT_RUBRIC_ID,
    }


_ANTHROPIC_MODEL = "claude-opus-4-7"

# Provider keys → actual Gemini model IDs. The two flavors have INDEPENDENT
# daily quotas on the Free Tier ("PerModel" in the rate-limit dimension), so
# offering both gives ~2× the free-tier headroom: 20 RPD each → 40 RPD total
# when the teacher manually flips models after the first hits its cap.
_GEMINI_MODEL_IDS: dict[str, str] = {
    "gemini": "gemini-2.5-flash",
    "gemini-lite": "gemini-2.5-flash-lite",
}


def _is_gemini(provider: str) -> bool:
    return provider in _GEMINI_MODEL_IDS


def _gemini_model_id(provider: str) -> str:
    return _GEMINI_MODEL_IDS.get(provider, "gemini-2.5-flash")

_gemini_client = None


def _gemini():
    """Lazily built so a missing GEMINI_API_KEY only bites if Gemini is
    actually picked (Claude-only deployments stay unaffected)."""
    global _gemini_client
    if _gemini_client is None:
        from google import genai

        _gemini_client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    return _gemini_client


# Provider keys → actual Groq model IDs. Free tier offers Llama 4 Scout with
# vision support at ~30 RPM. Quality on Hebrew handwriting is unproven vs.
# Claude/Gemini — kept as a fallback for when both quota out.
_GROQ_MODEL_IDS: dict[str, str] = {
    "groq-scout": "meta-llama/llama-4-scout-17b-16e-instruct",
}


def _is_groq(provider: str) -> bool:
    return provider in _GROQ_MODEL_IDS


def _groq_model_id(provider: str) -> str:
    return _GROQ_MODEL_IDS.get(provider, "meta-llama/llama-4-scout-17b-16e-instruct")


_groq_client = None


def _groq():
    """Lazily built so a missing GROQ_API_KEY only bites if Groq is actually
    picked."""
    global _groq_client
    if _groq_client is None:
        from groq import Groq

        _groq_client = Groq(api_key=os.environ["GROQ_API_KEY"])
    return _groq_client


def _ocr_anthropic(img: Image.Image) -> list[dict]:
    img_b64 = base64.standard_b64encode(ocr_jpeg_bytes(img)).decode("utf-8")
    response = client.messages.create(
        model=_ANTHROPIC_MODEL,
        max_tokens=8000,
        system=[
            {
                "type": "text",
                "text": OCR_SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        output_config={
            "format": {"type": "json_schema", "schema": OCR_SCHEMA}
        },
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": img_b64,
                        },
                    },
                    {"type": "text", "text": _USER_TURN},
                ],
            }
        ],
    )
    text = next(b.text for b in response.content if b.type == "text")
    return json.loads(text)["lines"]


def _ocr_gemini(img: Image.Image, provider: str = "gemini") -> list[dict]:
    from google.genai import types

    # Gemini's response_schema is an OpenAPI subset — it rejects the
    # additionalProperties key that OCR_SCHEMA carries for Anthropic.
    gemini_schema = {
        "type": "object",
        "properties": {
            "lines": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {"text": {"type": "string"}},
                    "required": ["text"],
                },
            }
        },
        "required": ["lines"],
    }
    response = _gemini().models.generate_content(
        model=_gemini_model_id(provider),
        contents=[
            types.Part.from_bytes(
                data=ocr_jpeg_bytes(img), mime_type="image/jpeg"
            ),
            _USER_TURN,
        ],
        config=types.GenerateContentConfig(
            system_instruction=OCR_SYSTEM_PROMPT,
            response_mime_type="application/json",
            response_schema=gemini_schema,
            # See evaluate path for why thinking is disabled here too.
            thinking_config=types.ThinkingConfig(thinking_budget=0),
            max_output_tokens=8000,
        ),
    )
    return json.loads(response.text)["lines"]


def _ocr_groq(img: Image.Image, provider: str = "groq-scout") -> list[dict]:
    # Groq is OpenAI-compatible; vision takes a data: URL in image_url. The
    # SDK only supports response_format={"type":"json_object"} (no schema),
    # so we lean on the prompt to enforce shape and parse defensively.
    img_b64 = base64.standard_b64encode(ocr_jpeg_bytes(img)).decode("utf-8")
    response = _groq().chat.completions.create(
        model=_groq_model_id(provider),
        messages=[
            {"role": "system", "content": OCR_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
                    },
                    {
                        "type": "text",
                        "text": (
                            _USER_TURN
                            + ' החזר JSON בפורמט הבא בלבד: {"lines":[{"text":"..."},...]}'
                        ),
                    },
                ],
            },
        ],
        response_format={"type": "json_object"},
        max_tokens=8000,
        temperature=0,
    )
    return json.loads(response.choices[0].message.content)["lines"]


def ocr_page(img: Image.Image, provider: str = DEFAULT_MODEL) -> list[dict]:
    if _is_gemini(provider):
        return _ocr_gemini(img, provider)
    if _is_groq(provider):
        return _ocr_groq(img, provider)
    return _ocr_anthropic(img)


def crop_line(img: Image.Image, y_top: int, y_bottom: int, padding: int = 0) -> str:
    h = img.height
    top = max(0, y_top - padding)
    bottom = min(h, y_bottom + padding)
    strip = img.crop((0, top, img.width, bottom))
    return image_to_b64(strip, "PNG")


def _smoothed_profile(
    img: Image.Image,
    ink_threshold: int = 180,
    smooth_window: int = 25,
    solid_row_ratio: float = 0.40,
    edge_margin_ratio: float = 0.04,
) -> np.ndarray:
    """Horizontal-projection ink profile, smoothed, with noise removed.

    solid_row_ratio: any row whose ink covers more than this fraction of the
        page width is zeroed out. Handwriting strokes are thin and gappy and
        never cover that much of a row — but a scan bar or heavy underline
        might.
    edge_margin_ratio: the top/bottom slice of the page is zeroed. A photo of
        a bound notebook almost always catches the spiral/clip and the desk
        shadow at the extreme top edge; its coverage is the *same* as a dense
        text row, so it can't be told apart by density — but real writing
        always has a page margin, so anything in the first/last few percent
        is noise. This was the root cause of the persistent off-by-one: the
        binding became a phantom first line and shifted every strip down one.
    """
    gray = np.asarray(img.convert("L"), dtype=np.uint8)
    h, w = gray.shape
    profile = (gray < ink_threshold).astype(np.float32).sum(axis=1)
    profile[profile > w * solid_row_ratio] = 0.0
    margin = int(h * edge_margin_ratio)
    if margin > 0:
        profile[:margin] = 0.0
        profile[h - margin:] = 0.0
    if smooth_window > 1:
        kernel = np.ones(smooth_window, dtype=np.float32) / smooth_window
        profile = np.convolve(profile, kernel, mode="same")
    return profile


def _axis_profile(img: Image.Image, axis: int) -> np.ndarray:
    """Smoothed ink profile along an axis. axis=1 sums per row, axis=0 per col.

    Used by auto_rotate to compare how 'peaky' the projection is along each
    axis. Unlike _smoothed_profile (which is tuned for upright pages and
    zeroes out solid rows), this keeps everything — notebook ruling lines
    appear as solid rows in one orientation and solid columns in the other,
    so they're the *cleanest* signal of which way the page is facing.
    """
    gray = np.asarray(img.convert("L"), dtype=np.uint8)
    ink = (gray < 180).astype(np.float32)
    profile = ink.sum(axis=axis)
    smooth_window = 25
    if smooth_window > 1:
        kernel = np.ones(smooth_window, dtype=np.float32) / smooth_window
        profile = np.convolve(profile, kernel, mode="same")
    return profile


def auto_rotate(img: Image.Image) -> tuple[Image.Image, int]:
    """Pick the 0/90/180/270° rotation that makes the text horizontal & top-heavy.

    EXIF-based fixup (exif_transpose) misses phones that don't write the
    Orientation tag and notebook spreads shot in landscape; deskew_page only
    handles a few degrees. This is the content-based fallback that catches
    those cases before the rest of the pipeline runs.

    Signal: compare std of the horizontal projection (sum per row) to std of
    the vertical projection (sum per column). Whichever axis the *line
    breaks* run perpendicular to has higher std — the projection alternates
    line→gap→line. So h_std > v_std ⇒ text lines stacked top-to-bottom (the
    image is already orient-ish); v_std > h_std ⇒ stacked side-to-side (the
    page is sideways and needs a 90° turn). Then ink-half ratios on the
    winning axis pick which 180° flip (top vs upside-down, or CW vs CCW).
    """
    # If EXIF already brought the photo to portrait, trust it. The heuristic
    # below gets fooled by background texture parallel to one axis (notebook
    # photographed on a towel whose ridges run vertically, etc.) — and the
    # case it was built for (notebook spread shot in landscape, phone with
    # no EXIF Orientation tag) still triggers because those come through
    # landscape.
    if img.height >= img.width:
        log.info("auto_rotate: image already portrait — skipping")
        return img, 0

    h_prof = _axis_profile(img, axis=1)
    v_prof = _axis_profile(img, axis=0)
    h_std = float(h_prof.std())
    v_std = float(v_prof.std())

    # Heuristic: the page header (date, name, title) lives at the TOP of a
    # school exercise — and it's *sparser* than the body of writing below
    # it. So the side we want to end up at the top is the side with LESS
    # ink. (Top-heavy = "upside-down body" or "header at bottom".)
    if h_std >= v_std:
        total = float(h_prof.sum()) or 1.0
        top_frac = float(h_prof[: len(h_prof) // 2].sum()) / total
        angle = 0 if top_frac <= 0.5 else 180
    else:
        total = float(v_prof.sum()) or 1.0
        left_frac = float(v_prof[: len(v_prof) // 2].sum()) / total
        # rotate(-90) = 90° CW: original LEFT edge becomes the new TOP.
        # rotate(-270) = 90° CCW: original RIGHT edge becomes the new TOP.
        # We want the lighter half to become the top.
        angle = 90 if left_frac <= 0.5 else 270

    rotated = (
        img
        if angle == 0
        else img.rotate(
            -angle,
            resample=Image.BILINEAR,
            expand=True,
            fillcolor=(255, 255, 255),
        )
    )
    log.info(
        "auto_rotate: h_std=%.1f v_std=%.1f → picked %d°",
        h_std, v_std, angle,
    )
    return rotated, angle


def deskew_page(img: Image.Image, max_angle: float = 6.0) -> tuple[Image.Image, float]:
    """Rotate the page so its text lines are horizontal.

    A phone photo of a notebook is almost always shot at a slight angle, so
    the ruled lines curve across the frame. A horizontal projection of a
    skewed page smears neighbouring lines into each other (shallow valleys —
    the root cause of merged/duplicated strips). We find the rotation that
    maximises profile contrast (std): when lines are level the projection has
    tall peaks and deep valleys, so its standard deviation peaks too.
    """
    def sharpness(angle: float) -> float:
        rot = img.rotate(
            angle, resample=Image.BILINEAR, fillcolor=(255, 255, 255)
        )
        return float(_smoothed_profile(rot).std())

    best_angle, best = 0.0, -1.0
    for a in np.arange(-max_angle, max_angle + 0.01, 0.5):
        s = sharpness(float(a))
        if s > best:
            best_angle, best = float(a), s
    for a in np.arange(best_angle - 0.5, best_angle + 0.51, 0.1):
        s = sharpness(float(a))
        if s > best:
            best_angle, best = float(a), s
    angle = round(best_angle, 1)
    if abs(angle) < 0.1:
        return img, 0.0
    return img.rotate(angle, resample=Image.BILINEAR, fillcolor=(255, 255, 255)), angle


def _writing_region(profile: np.ndarray) -> tuple[int, int]:
    """First/last row of *sustained* ink — a thin stray mark won't start it."""
    h = len(profile)
    if profile.max() <= 0:
        return 0, h
    present = profile > profile.max() * 0.10
    min_run = 30
    top, bottom = 0, h
    run = 0
    for y in range(h):
        if present[y]:
            run += 1
            if run >= min_run:
                top = y - run + 1
                break
        else:
            run = 0
    run = 0
    for y in range(h - 1, -1, -1):
        if present[y]:
            run += 1
            if run >= min_run:
                bottom = y + run
                break
        else:
            run = 0
    return top, min(h, bottom)


def _prominence(region: np.ndarray, i: int) -> float:
    """Topographic prominence of peak `i` — how far it rises above the
    deeper of the two valleys separating it from any taller neighbour."""
    left = i
    left_min = region[i]
    while left > 0 and region[left] <= region[i]:
        left_min = min(left_min, region[left])
        left -= 1
    right = i
    right_min = region[i]
    while right < len(region) - 1 and region[right] <= region[i]:
        right_min = min(right_min, region[right])
        right += 1
    return float(region[i] - max(left_min, right_min))


def segment_into_lines(img: Image.Image, num_lines: int) -> list[tuple[int, int]]:
    """Cut the writing region into exactly `num_lines` strips.

    We know the line count from the OCR, so rather than guess it from a
    threshold (which merges close lines and splits sparse ones), we place one
    peak per written line in the ink profile and cut at the lowest point
    between consecutive peaks.

    The first and last lines are *anchored* to the strongest bump in the
    first/last pitch-window; the middle `num_lines - 2` are the most
    prominent peaks in between, spaced at least half a pitch apart. Anchoring
    is what keeps a faint short opening line (a name/signature) or a faint
    closing line from being out-competed by the dense body lines — without
    it, a dense line grabs two peaks and a faint edge line gets none, which
    shifts a whole block of strips by one.
    """
    if num_lines <= 0:
        return []
    profile = _smoothed_profile(img)
    top, bottom = _writing_region(profile)
    if num_lines == 1 or bottom - top < num_lines:
        return [(top, bottom)]

    region = profile[top:bottom]
    floor = profile.max() * 0.05

    # Dominant line pitch (notebook ruling) → minimum spacing between peaks.
    # 0.5 (not higher): some lines are written consecutively with a tight gap,
    # so a larger guard would merge two real lines and lose one peak.
    centred = region - region.mean()
    autocorr = np.correlate(centred, centred, "full")[len(centred) - 1:]
    pitch = 80 + int(np.argmax(autocorr[80:220])) if len(autocorr) > 220 else 100
    pitch = min(pitch, len(region))
    min_dist = max(15, int(pitch * 0.5))

    candidates = [
        i
        for i in range(1, len(region) - 1)
        if region[i] >= region[i - 1] and region[i] > region[i + 1] and region[i] > floor
    ]

    if num_lines >= 2 and len(candidates) >= num_lines:
        first = int(np.argmax(region[:pitch]))
        last = len(region) - pitch + int(np.argmax(region[-pitch:]))
        middle = sorted(
            (c for c in candidates if first + min_dist <= c <= last - min_dist),
            key=lambda i: -_prominence(region, i),
        )
        peaks = [first, last]
        for c in middle:
            if all(abs(c - p) >= min_dist for p in peaks):
                peaks.append(c)
            if len(peaks) == num_lines:
                break
        peaks.sort()

    if num_lines < 2 or len(candidates) < num_lines or len(peaks) < num_lines:
        # Not enough structure — fall back to even spacing in the region.
        step = (bottom - top) / num_lines
        return [
            (int(round(top + k * step)), int(round(top + (k + 1) * step)))
            for k in range(num_lines)
        ]

    bounds = [0]
    for a, b in zip(peaks, peaks[1:]):
        bounds.append(a + int(np.argmin(region[a:b])))
    bounds.append(len(region) - 1)
    segments = [(top + bounds[i], top + bounds[i + 1]) for i in range(num_lines)]
    heights = [b - a for a, b in segments]
    ink = [int(profile[a:b].sum()) for a, b in segments]
    med_ink = sorted(ink)[len(ink) // 2]
    near_empty = [i + 1 for i, v in enumerate(ink) if v < 0.35 * med_ink]
    log.info(
        "segment: region=(%d,%d) lines=%d pitch=%d peaks=%d "
        "heights min/med/max=%d/%d/%d near_empty_strips=%s",
        top, bottom, num_lines, pitch, len(peaks),
        min(heights), sorted(heights)[len(heights) // 2], max(heights),
        near_empty or "none",
    )
    return segments


# Stage labels shown in the progress bar (one event per stage per page).
STAGE_LABELS = {
    "render": "מרנדר עמוד",
    "deskew": "מיישר הטיה",
    "ocr": "מתמלל",
    "segment": "מחתך לשורות",
    "crop": "חותך שורות",
}
_STEPS_PER_PAGE = 5  # render, deskew, ocr, segment, crop


def process_pdf_stream(file_bytes: bytes, ext: str, provider: str):
    """Run the OCR pipeline, yielding progress as it goes.

    Yields {"type": "progress", page, total_pages, stage, pct} after each
    pipeline stage (render → deskew → ocr → segment → crop), then a final
    {"type": "result", pages}. Lets the UI show a real determinate bar
    instead of a guess.
    """
    pages_imgs = file_to_page_images(file_bytes, ext)
    total = len(pages_imgs)
    total_steps = max(1, total * _STEPS_PER_PAGE)
    done = 0
    results = []

    def progress(stage: str, page: int):
        nonlocal done
        done += 1
        return {
            "type": "progress",
            "page": page,
            "total_pages": total,
            "stage": stage,
            "label": STAGE_LABELS.get(stage, stage),
            "pct": round(100 * done / total_steps),
        }

    for page_idx, img in enumerate(pages_imgs):
        p = page_idx + 1
        yield progress("render", p)
        img, rot_angle = auto_rotate(img)
        original_b64 = original_preview_b64(img)
        img, angle = deskew_page(img)
        yield progress("deskew", p)
        texts = [line["text"] for line in ocr_page(img, provider)]
        log.info(
            "[page %d] model=%s auto_rotate=%d° deskew=%s° %d lines transcribed",
            p, provider, rot_angle, angle, len(texts),
        )
        yield progress("ocr", p)
        bands = segment_into_lines(img, len(texts))
        yield progress("segment", p)
        line_items = [
            {"text": text, "image_b64": crop_line(img, y_top, y_bottom)}
            for text, (y_top, y_bottom) in zip(texts, bands)
        ]
        results.append(
            {"page": p, "lines": line_items, "original_b64": original_b64}
        )
        yield progress("crop", p)

    yield {"type": "result", "pages": results}


def humanize_error(exc: Exception, *, action: str = "הפעולה") -> dict:
    """Turn a low-level provider exception into a structured UI error.

    Returns a dict with keys the front-end renders directly:
      - message   : short Hebrew headline shown in the error banner
      - details   : longer Hebrew explanation + recommended next step
      - link_url  : optional URL the user can click to act (quota dash, key page)
      - link_text : Hebrew label for the link
      - technical : raw exception text for the collapsible "פרטים טכניים" block

    The "guess by string match" approach is deliberate: provider SDKs don't
    expose a clean enum for "this is a quota error vs an auth error", and the
    JSON payload is the most reliable signal we have.
    """
    raw = (str(exc) or repr(exc)).strip()
    technical = raw[:2000]

    # Gemini errors come from google.genai.errors. Imported lazily so a stripped-
    # down install without google-genai (Claude-only) still loads this module.
    try:
        from google.genai import errors as gerrors  # type: ignore
    except ImportError:
        gerrors = None  # noqa: N806

    # ---- Gemini: daily/per-minute quota ----
    if gerrors and isinstance(exc, gerrors.APIError) and getattr(exc, "code", None) == 429:
        quota_match = re.search(r"['\"]?quotaValue['\"]?\s*[:=]\s*['\"]?(\d+)", raw)
        per_day = "PerDay" in raw or "free_tier_requests" in raw
        quota_str = (
            f"{quota_match.group(1)} בקשות{' ליום' if per_day else ''} ב-Free Tier"
            if quota_match
            else "המכסה החינמית"
        )
        return {
            "message": f"חרגת ממכסת Gemini ({quota_str})",
            "details": (
                "המכסה היומית של Gemini ב-Free Tier מתאפסת בחצות שעון Pacific — "
                "כלומר ~10:00 בבוקר שעון ישראל. אפשרויות: לחכות לאיפוס, "
                "להחליף מודל ל-Claude בתפריט בראש הדף, או להוסיף billing ולעבור "
                "ל-Tier 1 (תקרה גבוהה משמעותית)."
            ),
            "link_url": "https://aistudio.google.com/rate-limit",
            "link_text": "צפה במכסה האישית שלך ב-AI Studio",
            "technical": technical,
        }

    # ---- Gemini: bad / missing API key ----
    if gerrors and isinstance(exc, gerrors.APIError) and getattr(exc, "code", None) in (401, 403):
        return {
            "message": "מפתח Gemini API לא תקף",
            "details": (
                "ודא שהמשתנה GEMINI_API_KEY מוגדר נכון (.env מקומית או Secret "
                "ב-Hugging Face Space) ושהמפתח עדיין פעיל."
            ),
            "link_url": "https://aistudio.google.com/apikey",
            "link_text": "נהל מפתחות ב-Google AI Studio",
            "technical": technical,
        }

    # ---- Gemini: server-side outage ----
    if gerrors and isinstance(exc, gerrors.ServerError):
        return {
            "message": "Gemini מחזיר שגיאת שרת",
            "details": "תקלה זמנית בצד של Google. נסה שוב בעוד דקה, או החלף ל-Claude.",
            "link_url": "https://status.cloud.google.com",
            "link_text": "סטטוס שירותי Google",
            "technical": technical,
        }

    # ---- Groq errors: SDK exposes them under groq module. Imported lazily so
    # a stripped-down install without groq (Claude/Gemini-only) still loads.
    try:
        import groq as _groq_mod  # type: ignore
    except ImportError:
        _groq_mod = None  # noqa: N806

    if _groq_mod and isinstance(exc, _groq_mod.RateLimitError):
        return {
            "message": "חרגת ממכסת Groq",
            "details": (
                "Groq Free Tier מגביל בקשות לדקה ולשעה. נסה שוב בעוד דקה, "
                "או החלף ל-Claude/Gemini בתפריט המודל."
            ),
            "link_url": "https://console.groq.com/settings/limits",
            "link_text": "צפה במכסה ב-Groq Console",
            "technical": technical,
        }

    if _groq_mod and isinstance(exc, _groq_mod.AuthenticationError):
        return {
            "message": "מפתח Groq API לא תקף",
            "details": (
                "ודא שהמשתנה GROQ_API_KEY מוגדר נכון (.env מקומית או Secret "
                "ב-Hugging Face Space) ושהמפתח עדיין פעיל."
            ),
            "link_url": "https://console.groq.com/keys",
            "link_text": "נהל מפתחות ב-Groq Console",
            "technical": technical,
        }

    if _groq_mod and isinstance(exc, (_groq_mod.APIConnectionError, _groq_mod.APITimeoutError)):
        return {
            "message": "החיבור ל-Groq נכשל",
            "details": "ייתכן שיש בעיית רשת או ש-Groq לא זמין כרגע. נסה שוב בעוד דקה.",
            "link_url": "https://groqstatus.com",
            "link_text": "סטטוס שירות Groq",
            "technical": technical,
        }

    # ---- Anthropic: rate limit ----
    if isinstance(exc, anthropic.RateLimitError):
        return {
            "message": "חרגת ממכסת Claude לרגע זה",
            "details": (
                "Anthropic מגביל בקשות לדקה (RPM) וטוקנים לדקה (TPM). "
                "נסה שוב בעוד דקה, או החלף ל-Gemini בתפריט המודל."
            ),
            "link_url": "https://console.anthropic.com/settings/limits",
            "link_text": "צפה במכסה ב-Anthropic Console",
            "technical": technical,
        }

    # ---- Anthropic: bad / missing API key ----
    if isinstance(exc, anthropic.AuthenticationError):
        return {
            "message": "מפתח Anthropic API לא תקף",
            "details": (
                "ודא שהמשתנה ANTHROPIC_API_KEY מוגדר נכון (.env מקומית או Secret "
                "ב-Hugging Face Space) ושהמפתח עדיין פעיל."
            ),
            "link_url": "https://console.anthropic.com/settings/keys",
            "link_text": "נהל מפתחות ב-Anthropic Console",
            "technical": technical,
        }

    # ---- Anthropic: network/timeout ----
    if isinstance(exc, (anthropic.APITimeoutError, anthropic.APIConnectionError)):
        return {
            "message": "החיבור ל-Anthropic נכשל",
            "details": "ייתכן שיש בעיית רשת או ש-Anthropic לא זמין כרגע. נסה שוב בעוד דקה.",
            "link_url": "https://status.anthropic.com",
            "link_text": "סטטוס שירות Anthropic",
            "technical": technical,
        }

    # ---- Missing API key env var: providers' lazy clients raise KeyError on
    # os.environ[...] when the key isn't set. Catch this BEFORE the generic
    # JSON branch below (which also matches KeyError) so the user sees the
    # real cause instead of "the model returned invalid response".
    if isinstance(exc, KeyError) and exc.args:
        key_name = str(exc.args[0])
        _key_info = {
            "GROQ_API_KEY": (
                "Groq",
                "https://console.groq.com/keys",
                "צור מפתח חינמי ב-Groq Console",
            ),
            "ANTHROPIC_API_KEY": (
                "Anthropic",
                "https://console.anthropic.com/settings/keys",
                "נהל מפתחות ב-Anthropic Console",
            ),
            "GEMINI_API_KEY": (
                "Gemini",
                "https://aistudio.google.com/apikey",
                "נהל מפתחות ב-Google AI Studio",
            ),
        }
        if key_name in _key_info:
            provider_name, url, link_label = _key_info[key_name]
            return {
                "message": f"מפתח {provider_name} API חסר",
                "details": (
                    f"המשתנה {key_name} לא מוגדר. הוסף אותו ל-.env המקומי או "
                    f"כ-Secret ב-Hugging Face Space, ואז restart את השרת."
                ),
                "link_url": url,
                "link_text": link_label,
                "technical": technical,
            }

    # ---- JSON parsing: the model returned malformed output ----
    if isinstance(exc, (json.JSONDecodeError, KeyError, ValueError)):
        return {
            "message": "המודל החזיר תשובה לא תקפה",
            "details": (
                f"לא הצלחנו לפענח את התשובה של המודל ({action}). "
                "זו בדרך כלל בעיה זמנית — נסה שוב, ואם זה נמשך, החלף מודל."
            ),
            "technical": technical,
        }

    # ---- Generic provider 5xx (when not caught above) ----
    if isinstance(exc, anthropic.APIStatusError):
        return {
            "message": f"שגיאת API ({getattr(exc, 'status_code', '?')})",
            "details": "השירות החזיר תשובה לא תקינה. נסה שוב, ואם זה נמשך החלף מודל.",
            "technical": technical,
        }

    # ---- Fallback ----
    return {
        "message": "תקלה לא צפויה",
        "details": "אירעה שגיאה לא מזוהה. נסה שוב; אם זה חוזר, צרף את הפרטים הטכניים בדיווח.",
        "technical": technical,
    }


# In-memory job store. Single local user (Avishai), so a dict is enough; a
# job is short-lived and consumed once by the SSE stream.
JOBS: dict[str, dict] = {}


def run_job(job_id: str, file_bytes: bytes, ext: str, provider: str):
    job = JOBS[job_id]
    q: queue.Queue = job["q"]
    try:
        for ev in process_pdf_stream(file_bytes, ext, provider):
            if ev["type"] == "result":
                job["pages"] = ev["pages"]
                q.put({"type": "done"})
            else:
                q.put(ev)
    except Exception as e:
        # Catch-all so a worker-thread crash surfaces in the SSE stream
        # instead of leaving the UI stuck on "מתחיל…". humanize_error sorts
        # quota / auth / parse / unknown into user-readable messages.
        log.exception("run_job failed")
        q.put({"type": "error", **humanize_error(e, action="הפענוח")})
    finally:
        q.put({"type": "end"})


def compute_doc_key(pages: list[dict]) -> str:
    """Stable hash of the line texts — used as localStorage key."""
    joined = "\n".join(
        line["text"]
        for page in pages
        for line in page.get("lines", [])
    )
    return hashlib.sha1(joined.encode("utf-8")).hexdigest()[:16]


def render_result(
    pages: list[dict],
    filename: str,
    annotations: dict | None = None,
    evaluation: dict | None = None,
):
    server_data = {
        "doc_key": compute_doc_key(pages),
        "filename": filename,
        "annotations": annotations or {},
        "evaluation": evaluation,
        "highlights": _highlights_for_template(pages, evaluation),
        "rubrics": list_rubrics(),
        "default_rubric_id": DEFAULT_RUBRIC_ID,
    }
    return render_template(
        "index.html",
        pages=pages,
        filename=filename,
        server_data_json=json.dumps(server_data, ensure_ascii=False),
    )


def _highlights_for_template(
    pages: list[dict], evaluation: dict | None
) -> dict[str, list[dict]]:
    """Convert resolve_issue_spans output to the JSON shape the browser
    JS consumes: { "p1-l3": [{start, end, color, comment, criterion}] }."""
    if not evaluation:
        return {}
    raw = resolve_issue_spans(pages, attach_colors(evaluation))
    return {f"p{page}-l{idx}": spans for (page, idx), spans in raw.items()}


def parse_saved_json(raw: bytes) -> tuple[list[dict], str, dict, dict | None]:
    """Returns (pages_for_template, filename, annotations_by_line_key, evaluation).

    Strips annotations out of the per-line dicts (the template doesn't need
    them inline — they're delivered via server_data_json instead).
    `evaluation` is the saved rubric scoring (or None if the file predates
    the feature or no evaluation was run).
    """
    data = json.loads(raw.decode("utf-8"))
    if not isinstance(data, dict) or "pages" not in data:
        raise ValueError("מבנה JSON לא תקין — חסר שדה 'pages'")

    pages_out = []
    annotations: dict[str, list] = {}
    for page in data["pages"]:
        page_num = page.get("page", len(pages_out) + 1)
        lines_out = []
        for line_idx, line in enumerate(page.get("lines", [])):
            text = line.get("text", "")
            image_b64 = line.get("image_b64", "")
            line_key = f"p{page_num}-l{line_idx}"
            line_annots = line.get("annotations") or []
            if line_annots:
                annotations[line_key] = line_annots
            lines_out.append({"text": text, "image_b64": image_b64})
        pages_out.append({
            "page": page_num,
            "lines": lines_out,
            "original_b64": page.get("original_b64", ""),
        })

    filename = data.get("filename", "שמור.json")
    evaluation = attach_colors(data.get("evaluation"))
    return pages_out, filename, annotations, evaluation


# --- Rubrics ----------------------------------------------------------------

RUBRICS_DIR = os.path.join(os.path.dirname(__file__), "rubrics")

# Initial rubric the dropdown should land on for a fresh user (no
# localStorage). Any previous explicit choice still wins on subsequent
# loads — this only affects "first visit on this browser" / cleared state.
DEFAULT_RUBRIC_ID = "ministry-writing-module-c-d"


def _slugify(name: str) -> str:
    """Filename-safe slug. Keeps Hebrew/Arabic letters; replaces whitespace
    and forbidden filesystem chars with '-'."""
    s = re.sub(r"[\s/\\:*?\"<>|]+", "-", name.strip())
    s = re.sub(r"-+", "-", s).strip("-")
    return s or "rubric"


def list_rubrics() -> list[dict]:
    """[{id, name}] for every .json file in rubrics/. id == filename stem."""
    if not os.path.isdir(RUBRICS_DIR):
        return []
    out = []
    for fname in sorted(os.listdir(RUBRICS_DIR)):
        if not fname.endswith(".json"):
            continue
        path = os.path.join(RUBRICS_DIR, fname)
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            out.append({"id": fname[:-5], "name": data.get("name", fname[:-5])})
        except (json.JSONDecodeError, OSError):
            continue
    return out


def load_rubric(rubric_id: str) -> dict | None:
    """Returns {name, content} or None if not found."""
    safe = _slugify(rubric_id)
    path = os.path.join(RUBRICS_DIR, f"{safe}.json")
    if not os.path.isfile(path):
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return None


# --- Evaluation -------------------------------------------------------------

EVAL_SYSTEM_PROMPT = """אתה בודק שיעורי בית בהתאם לרובריקה שתינתן לך.

תקבל:
1. רובריקה — מתארת קריטריונים, רמות וציונים אפשריים.
2. טקסט תרגיל של תלמיד — תמלול של כתב יד, יתכן עם שגיאות OCR. כל
   שורה מתויגת עם מזהה בצורה [p<page>-l<index>] בתחילתה.

המשימה:
- זהה את כל הקריטריונים שמופיעים ברובריקה (השמות שלהם, והציון המקסימלי בכל אחד).
- העריך את התרגיל לפי הרובריקה — צא מנקודת הנחה שמה שנכתב הוא מה שהתלמיד התכוון אליו (אל תקטף נקודות על שגיאות OCR שנראות כמו טעויות תמלול).
- לכל קריטריון: ציון מספרי (1 עד max_score לפי הרובריקה) ופידבק קצר בעברית (1-3 משפטים).
- לכל קריטריון: גם פידבק באנגלית ("feedback_en") — תרגום נאמן של הפידבק
  העברי, באותו אורך ובאותו תוכן (לא להוסיף או להחסיר מידע).
- לכל קריטריון: רשימת "issues" — דוגמאות ספציפיות לבעיות שמצאת
  בטקסט. לכל issue:
    * "line_ref": המזהה של השורה שבה הבעיה (לדוגמה "p1-l3").
    * "quote": ציטוט מילולי מדויק של הקטע הבעייתי מתוך אותה שורה —
      בדיוק כפי שהוא מופיע בתמלול (אותו רישיות, פיסוק ורווחים). זה
      חייב להיות תת-מחרוזת של השורה. ציטוט ברמת משפט/ביטוי, לא
      מילה בודדת ולא שורה שלמה אם רק חלק ממנה בעייתי.
    * "comment": הערה קצרה בעברית (3-15 מילים) שמסבירה למה זה בעייתי.
  אם אין בעיות בקריטריון מסוים, החזר רשימה ריקה.
  אם אותו משפט בעייתי בכמה קריטריונים — שייך אותו לקריטריון
  החמור/המהותי ביותר בלבד (לא לכמה במקביל).
- ציון כללי (אם הרובריקה מציינת mapping ל-CEFR/אחוז/אות — השתמש בו, אחרת תן את ממוצע הציונים).
- פסקת סיכום בעברית (2-4 משפטים) — חוזקות, חולשות, ומה כדאי לתלמיד לעבוד עליו.
- גם פסקת סיכום באנגלית ("overall_feedback_en") — תרגום נאמן של פסקת
  הסיכום העברית, באותו אורך ובאותו תוכן.

החזר JSON תקין במבנה:
{
  "criteria": [
    {
      "name": "...",
      "score": <int>,
      "max_score": <int>,
      "feedback": "...",
      "feedback_en": "...",
      "issues": [
        {"line_ref": "p1-l3", "quote": "...", "comment": "..."}
      ]
    }
  ],
  "overall_score": "<string — לדוגמה: B1, או 3.0/4, או 75%>",
  "overall_feedback": "...",
  "overall_feedback_en": "..."
}"""

EVAL_SCHEMA = {
    "type": "object",
    "properties": {
        "criteria": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "score": {"type": "number"},
                    "max_score": {"type": "number"},
                    "feedback": {"type": "string"},
                    "feedback_en": {"type": "string"},
                    "issues": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "line_ref": {"type": "string"},
                                "quote": {"type": "string"},
                                "comment": {"type": "string"},
                            },
                            "required": ["line_ref", "quote", "comment"],
                            "additionalProperties": False,
                        },
                    },
                },
                "required": ["name", "score", "max_score", "feedback", "feedback_en", "issues"],
                "additionalProperties": False,
            },
        },
        "overall_score": {"type": "string"},
        "overall_feedback": {"type": "string"},
        "overall_feedback_en": {"type": "string"},
    },
    "required": ["criteria", "overall_score", "overall_feedback", "overall_feedback_en"],
    "additionalProperties": False,
}


# Groq's chat-completions API supports response_format={"type":"json_object"}
# but NOT json_schema, so we spell the required shape in the prompt instead.
_GROQ_EVAL_SHAPE_HINT = (
    'החזר אך ורק JSON תקין בפורמט הבא:\n'
    '{\n'
    '  "criteria": [\n'
    '    {\n'
    '      "name": "...",\n'
    '      "score": 0,\n'
    '      "max_score": 0,\n'
    '      "feedback": "...",\n'
    '      "feedback_en": "...",\n'
    '      "issues": [\n'
    '        {"line_ref": "p1-l3", "quote": "...", "comment": "..."}\n'
    '      ]\n'
    '    }\n'
    '  ],\n'
    '  "overall_score": "...",\n'
    '  "overall_feedback": "...",\n'
    '  "overall_feedback_en": "..."\n'
    '}'
)


def pages_to_plain_text(pages: list[dict]) -> str:
    """Flatten the OCR'd pages into a single transcript for the evaluator.

    Each line is prefixed with its [p<page>-l<idx>] tag so the model can
    return per-issue references that we map back to highlights later."""
    chunks = []
    for page in pages:
        page_num = page.get("page", "?")
        chunks.append(f"--- עמוד {page_num} ---")
        for line_idx, line in enumerate(page.get("lines", [])):
            chunks.append(f"[p{page_num}-l{line_idx}] {line.get('text', '')}")
        chunks.append("")
    return "\n".join(chunks).strip()


def evaluate_with_rubric(pages: list[dict], rubric: dict, provider: str) -> dict:
    """Send transcript + rubric → structured per-criterion scores + feedback."""
    transcript = pages_to_plain_text(pages)
    user_turn = (
        f"רובריקה (שם: {rubric.get('name', '')}):\n\n"
        f"{rubric.get('content', '')}\n\n"
        f"--- טקסט התרגיל לבדיקה ---\n\n{transcript}\n\n"
        "החזר JSON לפי הסכימה."
    )
    if _is_gemini(provider):
        from google.genai import types

        gemini_schema = {
            "type": "object",
            "properties": {
                "criteria": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "score": {"type": "number"},
                            "max_score": {"type": "number"},
                            "feedback": {"type": "string"},
                            "feedback_en": {"type": "string"},
                            "issues": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "line_ref": {"type": "string"},
                                        "quote": {"type": "string"},
                                        "comment": {"type": "string"},
                                    },
                                    "required": ["line_ref", "quote", "comment"],
                                },
                            },
                        },
                        "required": ["name", "score", "max_score", "feedback", "feedback_en", "issues"],
                    },
                },
                "overall_score": {"type": "string"},
                "overall_feedback": {"type": "string"},
                "overall_feedback_en": {"type": "string"},
            },
            "required": ["criteria", "overall_score", "overall_feedback", "overall_feedback_en"],
        }
        response = _gemini().models.generate_content(
            model=_gemini_model_id(provider),
            contents=[user_turn],
            config=types.GenerateContentConfig(
                system_instruction=EVAL_SYSTEM_PROMPT,
                response_mime_type="application/json",
                response_schema=gemini_schema,
                # Gemini 2.5 Flash is a thinking model — without this it
                # silently eats most of max_output_tokens on internal
                # reasoning and the JSON gets truncated mid-string.
                thinking_config=types.ThinkingConfig(thinking_budget=0),
                max_output_tokens=12000,
            ),
        )
        try:
            parsed = json.loads(response.text)
        except json.JSONDecodeError as e:
            log.error(
                "evaluate (gemini) JSON parse failed: %s | text_len=%d",
                e, len(response.text),
            )
            log.error("evaluate (gemini) raw response text:\n%s", response.text)
            raise
        return attach_colors(parsed)

    if _is_groq(provider):
        # Groq's response_format only supports json_object (no schema), so we
        # spell the shape out in the user turn and parse defensively.
        response = _groq().chat.completions.create(
            model=_groq_model_id(provider),
            messages=[
                {"role": "system", "content": EVAL_SYSTEM_PROMPT},
                {"role": "user", "content": user_turn + "\n\n" + _GROQ_EVAL_SHAPE_HINT},
            ],
            response_format={"type": "json_object"},
            max_tokens=12000,
            temperature=0,
        )
        text = response.choices[0].message.content
        try:
            parsed = json.loads(text)
        except json.JSONDecodeError as e:
            log.error(
                "evaluate (groq) JSON parse failed: %s | text_len=%d", e, len(text),
            )
            log.error("evaluate (groq) raw response text:\n%s", text)
            raise
        return attach_colors(parsed)

    response = client.messages.create(
        model=_ANTHROPIC_MODEL,
        max_tokens=12000,
        system=[
            {
                "type": "text",
                "text": EVAL_SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        output_config={
            "format": {"type": "json_schema", "schema": EVAL_SCHEMA}
        },
        messages=[{"role": "user", "content": user_turn}],
    )
    text = next(b.text for b in response.content if b.type == "text")
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as e:
        stop = getattr(response, "stop_reason", "?")
        usage = getattr(response, "usage", None)
        log.error(
            "evaluate JSON parse failed: %s | stop_reason=%s | usage=%s | text_len=%d",
            e, stop, usage, len(text),
        )
        log.error("evaluate raw response text:\n%s", text)
        raise
    return attach_colors(parsed)


# Matches the criterion `"name": "..."` field in the streaming JSON. Issues
# don't have a `name` field (line_ref/quote/comment), so this only fires at
# the criterion level. The escape-aware pattern survives names that contain
# an escaped quote without prematurely terminating the match.
_CRITERION_NAME_RE = re.compile(r'"name"\s*:\s*"((?:[^"\\]|\\.)*)"')

# Used as a soft hint for "N/total" in progress messages. Both bundled
# rubrics mark criteria with "CRITERION N" / "Criterion N"; a user rubric
# that doesn't follow this convention just gets "N" without total.
_RUBRIC_CRITERION_RE = re.compile(r"(?i)\bcriterion\s+\d+\b")


def count_rubric_criteria(content: str) -> int:
    return len(_RUBRIC_CRITERION_RE.findall(content))


def evaluate_with_rubric_stream(pages: list[dict], rubric: dict, provider: str):
    """Streaming variant of evaluate_with_rubric.

    Yields:
      {"type": "progress", "index": N, "total": T|0, "name": "..."} as each
        criterion name is received (lets the UI say "בודק 2/4: VOCABULARY").
      {"type": "result", "evaluation": {...}} once the stream completes and
        the full JSON parses.

    On a malformed JSON at the end, logs the raw text and re-raises the
    JSONDecodeError so the worker emits a normal error event.
    """
    transcript = pages_to_plain_text(pages)
    user_turn = (
        f"רובריקה (שם: {rubric.get('name', '')}):\n\n"
        f"{rubric.get('content', '')}\n\n"
        f"--- טקסט התרגיל לבדיקה ---\n\n{transcript}\n\n"
        "החזר JSON לפי הסכימה."
    )
    total = count_rubric_criteria(rubric.get("content", ""))
    buffer = ""
    seen = 0

    def drain_progress():
        nonlocal seen
        names = _CRITERION_NAME_RE.findall(buffer)
        events = []
        for i in range(seen, len(names)):
            events.append({
                "type": "progress",
                "index": i + 1,
                "total": total,
                "name": names[i],
            })
        seen = len(names)
        return events

    if _is_gemini(provider):
        from google.genai import types

        gemini_schema = {
            "type": "object",
            "properties": {
                "criteria": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "score": {"type": "number"},
                            "max_score": {"type": "number"},
                            "feedback": {"type": "string"},
                            "feedback_en": {"type": "string"},
                            "issues": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "line_ref": {"type": "string"},
                                        "quote": {"type": "string"},
                                        "comment": {"type": "string"},
                                    },
                                    "required": ["line_ref", "quote", "comment"],
                                },
                            },
                        },
                        "required": ["name", "score", "max_score", "feedback", "feedback_en", "issues"],
                    },
                },
                "overall_score": {"type": "string"},
                "overall_feedback": {"type": "string"},
                "overall_feedback_en": {"type": "string"},
            },
            "required": ["criteria", "overall_score", "overall_feedback", "overall_feedback_en"],
        }
        stream = _gemini().models.generate_content_stream(
            model=_gemini_model_id(provider),
            contents=[user_turn],
            config=types.GenerateContentConfig(
                system_instruction=EVAL_SYSTEM_PROMPT,
                response_mime_type="application/json",
                response_schema=gemini_schema,
                thinking_config=types.ThinkingConfig(thinking_budget=0),
                max_output_tokens=12000,
            ),
        )
        for chunk in stream:
            if chunk.text:
                buffer += chunk.text
                for ev in drain_progress():
                    yield ev
    elif _is_groq(provider):
        stream = _groq().chat.completions.create(
            model=_groq_model_id(provider),
            messages=[
                {"role": "system", "content": EVAL_SYSTEM_PROMPT},
                {"role": "user", "content": user_turn + "\n\n" + _GROQ_EVAL_SHAPE_HINT},
            ],
            response_format={"type": "json_object"},
            max_tokens=12000,
            temperature=0,
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content if chunk.choices else None
            if delta:
                buffer += delta
                for ev in drain_progress():
                    yield ev
    else:
        with client.messages.stream(
            model=_ANTHROPIC_MODEL,
            max_tokens=12000,
            system=[
                {
                    "type": "text",
                    "text": EVAL_SYSTEM_PROMPT,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            output_config={
                "format": {"type": "json_schema", "schema": EVAL_SCHEMA}
            },
            messages=[{"role": "user", "content": user_turn}],
        ) as stream:
            for text_delta in stream.text_stream:
                buffer += text_delta
                for ev in drain_progress():
                    yield ev

    try:
        parsed = json.loads(buffer)
    except json.JSONDecodeError as e:
        log.error(
            "evaluate (%s, streaming) JSON parse failed: %s | text_len=%d",
            provider, e, len(buffer),
        )
        log.error("evaluate (%s, streaming) raw text:\n%s", provider, buffer)
        raise

    yield {"type": "result", "evaluation": attach_colors(parsed)}


# --- Criterion colors -------------------------------------------------------

# Fixed palette of 11 highly-distinct colors. Slots 0-3 are reserved for the
# default ministry-rubric criteria (in descending frequency of inline marks
# per teacher experience), so they never collide with anything else. Slots
# 4-10 absorb every other criterion via md5 → so the same criterion name
# always gets the same color across runs (teachers can train on the mapping).
_PALETTE: list[str] = [
    "#d9534f",  # 0 red
    "#2c7be5",  # 1 blue
    "#28a745",  # 2 green
    "#f0a000",  # 3 amber (richer than #e0a800 so it reads on cream paper)
    "#8e44ad",  # 4 purple
    "#e67e22",  # 5 orange
    "#16a085",  # 6 teal
    "#d63384",  # 7 magenta
    "#8B4513",  # 8 brown
    "#1a237e",  # 9 navy
    "#689f38",  # 10 lime
]

# Top-of-palette pins for the ministry rubric (the current default). Order
# from most-frequent inline marks (mechanics: every spelling slip) to least
# (content: a few essay-level notes). Keyword lists tolerate the LLM
# returning slight renames (e.g. "Content & Organization", or a Hebrew
# title). To pin a new rubric's criteria, append a tuple here.
_PINNED_CRITERION_RULES: list[tuple[list[str], int]] = [
    # Slot 0 (red) — Mechanics / spelling / punctuation
    (["mechanics", "spelling", "punctuation",
      "איות", "כתיב", "פיסוק"], 0),
    # Slot 1 (blue) — Language Use / grammar / syntax
    (["language use", "grammar", "syntax",
      "תחביר", "דקדוק"], 1),
    # Slot 2 (green) — Vocabulary / lexical
    (["vocabulary", "lexical", "lexicon", "word choice",
      "אוצר מילים"], 2),
    # Slot 3 (amber) — Content / Organization / ideas / structure
    (["content", "organization", "structure", "ideas",
      "תוכן", "ארגון", "מבנה", "רעיונות"], 3),
]

# md5 fallback only picks from slots 4..end, so a non-pinned criterion can
# never accidentally land on a pinned color.
_FALLBACK_RANGE = list(range(4, len(_PALETTE)))


def color_for_criterion(name: str) -> str:
    """Hex color for a rubric criterion. Pinned slots first (so the four
    ministry-rubric criteria always get red/blue/green/amber), else a
    deterministic md5 → palette slot from the remaining colors so the same
    criterion name always lands on the same color across runs."""
    n = (name or "").lower()
    for keywords, slot in _PINNED_CRITERION_RULES:
        if any(k.lower() in n for k in keywords):
            return _PALETTE[slot]
    digest = hashlib.md5(n.encode("utf-8")).hexdigest()
    slot = _FALLBACK_RANGE[int(digest, 16) % len(_FALLBACK_RANGE)]
    return _PALETTE[slot]


def attach_colors(evaluation: dict | None) -> dict | None:
    """Re-derive the 'color' field on every criterion from its name. We do
    NOT preserve a pre-existing color: the fixed palette is the source of
    truth, so loading an older saved JSON should still surface the current
    color scheme (teachers train on a stable red/blue/green/amber)."""
    if not evaluation:
        return evaluation
    for c in evaluation.get("criteria") or []:
        c["color"] = color_for_criterion(c.get("name", ""))
    return evaluation


# --- Issue span resolution --------------------------------------------------

_LINE_REF_RE = re.compile(r"p(\d+)-l(\d+)")
# Fuzzy fallback only forgives near-exact drift (punctuation, trailing space,
# 1-2 OCR slips). The model is instructed to return verbatim quotes, so we
# don't want to accept loose matches that paint the wrong span.
_FUZZY_MIN_BLOCK_RATIO = 0.8


def _find_span(line_text: str, quote: str) -> tuple[int, int] | None:
    """Locate `quote` within `line_text`. Tries exact, then case-insensitive,
    then a difflib longest-block fallback for small punctuation drift.
    Returns (start, end) on the original string, or None if no good match."""
    if not quote or not line_text:
        return None
    idx = line_text.find(quote)
    if idx >= 0:
        return idx, idx + len(quote)
    lower_line = line_text.lower()
    lower_quote = quote.lower()
    idx = lower_line.find(lower_quote)
    if idx >= 0:
        return idx, idx + len(lower_quote)
    matcher = difflib.SequenceMatcher(None, line_text, quote, autojunk=False)
    block = matcher.find_longest_match(0, len(line_text), 0, len(quote))
    min_size = max(4, int(len(quote) * _FUZZY_MIN_BLOCK_RATIO))
    if block.size < min_size:
        return None
    return block.a, block.a + block.size


def resolve_issue_spans(
    pages: list[dict], evaluation: dict | None
) -> dict[tuple[int, int], list[dict]]:
    """Map every (page, line_idx) to a list of highlight spans built from
    the evaluation's per-criterion issues. First criterion wins on overlap
    — a span that intersects an already-marked range is dropped (per the
    product decision: one color per chunk of text)."""
    if not evaluation:
        return {}
    spans: dict[tuple[int, int], list[dict]] = {}
    line_text_by_key: dict[tuple[int, int], str] = {}
    for page in pages:
        page_num = int(page.get("page", 0)) if str(page.get("page", "")).isdigit() else page.get("page", 0)
        for line_idx, line in enumerate(page.get("lines", [])):
            line_text_by_key[(int(page_num), line_idx)] = line.get("text", "")

    for crit in evaluation.get("criteria") or []:
        color = crit.get("color") or color_for_criterion(crit.get("name", ""))
        crit_name = crit.get("name", "")
        for issue in crit.get("issues") or []:
            ref = issue.get("line_ref", "")
            m = _LINE_REF_RE.search(ref)
            if not m:
                log.warning("issue with unparseable line_ref %r — skipped", ref)
                continue
            key = (int(m.group(1)), int(m.group(2)))
            line_text = line_text_by_key.get(key)
            if line_text is None:
                log.warning("issue references missing line %s — skipped", ref)
                continue
            quote = issue.get("quote", "")
            span = _find_span(line_text, quote)
            if not span:
                log.warning(
                    "issue quote not found in %s: %r — skipped", ref, quote[:50]
                )
                continue
            start, end = span
            existing = spans.setdefault(key, [])
            if any(not (end <= s["start"] or start >= s["end"]) for s in existing):
                continue  # first criterion wins on overlap
            existing.append(
                {
                    "start": start,
                    "end": end,
                    "color": color,
                    "comment": issue.get("comment", ""),
                    "criterion": crit_name,
                }
            )

    for key in spans:
        spans[key].sort(key=lambda s: s["start"])
    return spans


# --- DOCX export ------------------------------------------------------------

def _tint_hex(hex_color: str, mix: float = 0.18) -> str:
    """Blend hex_color toward white. mix=0.18 → ~82% white tint, light
    enough to keep black text readable on top."""
    c = hex_color.lstrip("#")
    r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
    r = int(r * mix + 255 * (1 - mix))
    g = int(g * mix + 255 * (1 - mix))
    b = int(b * mix + 255 * (1 - mix))
    return f"{r:02x}{g:02x}{b:02x}"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Fill a table cell with a background color. python-docx exposes no
    helper for this, so we drop to the OOXML layer."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_run_shading(run, hex_color: str) -> None:
    """Background shade a single run (used for highlighting a substring
    inside a paragraph). python-docx has no API for run-level shading."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    rPr.append(shd)


def _fill_paragraph_with_spans(
    paragraph, text: str, spans: list[dict], include_notes: bool = True
) -> None:
    """Empty `paragraph`, then add runs reflecting `spans` as colored
    highlights over `text`. When `include_notes` is True, append an
    inline `(criterion: comment)` run shaded in the same color after
    every highlighted span. Spans are assumed sorted, non-overlapping."""
    from docx.shared import Pt

    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)
    if not text:
        return
    cursor = 0
    for span in spans:
        s, e = span["start"], span["end"]
        if s > cursor:
            paragraph.add_run(text[cursor:s])
        tint = _tint_hex(span["color"], mix=0.35)
        run = paragraph.add_run(text[s:e])
        _set_run_shading(run, tint)
        if include_notes:
            comment = span.get("comment", "")
            criterion = span.get("criterion", "")
            if comment or criterion:
                note = (
                    f" ({criterion}: {comment})"
                    if criterion and comment
                    else f" ({criterion or comment})"
                )
                note_run = paragraph.add_run(note)
                note_run.italic = True
                note_run.font.size = Pt(9)
                _set_run_shading(note_run, tint)
        cursor = e
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def build_evaluation_docx(
    evaluation: dict,
    filename: str,
    rubric_name: str,
    pages: list[dict] | None = None,
) -> bytes:
    """Word document with the evaluation table + overall feedback. When
    `pages` is provided, appends each page's original scan followed by a
    line-by-line table (cropped line image | transcribed text) so the
    teacher can verify the OCR alongside the rubric scoring.

    Hebrew-friendly: paragraphs aligned right; the document is built top-down
    so existing readers (Google Docs, Word, LibreOffice) all open it cleanly.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()

    title = doc.add_heading("בדיקת תרגיל — השכלה", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"קובץ: {filename}\n").bold = True
    meta.add_run(f"רובריקה: {rubric_name}\n").bold = True
    meta.add_run(f"ציון כללי: {evaluation.get('overall_score', '')}").bold = True

    doc.add_heading("פירוט לפי קריטריון", level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    criteria = attach_colors(evaluation).get("criteria") or []
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "קריטריון"
    header[1].text = "ציון"
    header[2].text = "פידבק"
    header[3].text = "Feedback (EN)"
    for idx, cell in enumerate(header):
        align = WD_ALIGN_PARAGRAPH.LEFT if idx == 3 else WD_ALIGN_PARAGRAPH.RIGHT
        for p in cell.paragraphs:
            p.alignment = align
            for r in p.runs:
                r.bold = True

    for c in criteria:
        row = table.add_row().cells
        row[0].text = str(c.get("name", ""))
        row[1].text = f"{c.get('score', '')}/{c.get('max_score', '')}"
        row[2].text = str(c.get("feedback", ""))
        row[3].text = str(c.get("feedback_en", ""))
        for idx, cell in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 3 else WD_ALIGN_PARAGRAPH.RIGHT
            for p in cell.paragraphs:
                p.alignment = align
        color = c.get("color") or color_for_criterion(c.get("name", ""))
        _set_cell_shading(row[0], _tint_hex(color))
        for p in row[0].paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_heading("סיכום", level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary = doc.add_paragraph(evaluation.get("overall_feedback", ""))
    summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in summary.runs:
        r.font.size = Pt(11)

    overall_en = evaluation.get("overall_feedback_en", "")
    if overall_en:
        doc.add_heading("Summary", level=2).alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary_en = doc.add_paragraph(overall_en)
        summary_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in summary_en.runs:
            r.font.size = Pt(11)

    if pages:
        doc.add_page_break()
        spans_by_line = resolve_issue_spans(pages, evaluation)

        def _line_key(page_num, line_idx):
            try:
                return (int(page_num), line_idx)
            except (TypeError, ValueError):
                return (page_num, line_idx)

        # Section A: clean transcript — exactly what the student wrote, no
        # marks, so the teacher can read the raw text without distractions.
        # English content is left-aligned per the way English text reads.
        doc.add_heading("תעתיק נקי", level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for page in pages:
            page_num = page.get("page", "?")
            doc.add_heading(f"עמוד {page_num}", level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for line in page.get("lines", []):
                p = doc.add_paragraph(str(line.get("text", "")))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Section B: same text with the color highlights, but without the
        # inline `(criterion: comment)` parentheticals — for a scannable
        # at-a-glance view of where the issues fall.
        doc.add_heading("תעתיק עם סימונים", level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for page in pages:
            page_num = page.get("page", "?")
            doc.add_heading(f"עמוד {page_num}", level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for line_idx, line in enumerate(page.get("lines", [])):
                line_text = str(line.get("text", ""))
                spans = spans_by_line.get(_line_key(page_num, line_idx), [])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if spans:
                    _fill_paragraph_with_spans(p, line_text, spans, include_notes=False)
                else:
                    p.add_run(line_text)

        # Section C: the original scan + line-by-line table with image and
        # transcribed text (with both highlights and inline comments).
        doc.add_heading("תרגיל מקורי ופענוח", level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for page in pages:
            page_num = page.get("page", "?")
            doc.add_heading(f"עמוד {page_num}", level=3).alignment = WD_ALIGN_PARAGRAPH.RIGHT

            orig_b64 = page.get("original_b64") or ""
            if orig_b64:
                try:
                    orig_bytes = base64.b64decode(orig_b64)
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_para.add_run().add_picture(
                        io.BytesIO(orig_bytes), width=Inches(5.5)
                    )
                except (ValueError, OSError) as e:
                    log.warning("page %s original image failed: %s", page_num, e)

            lines = page.get("lines") or []
            if not lines:
                continue
            ltable = doc.add_table(rows=1, cols=2)
            ltable.style = "Light Grid Accent 1"
            lhdr = ltable.rows[0].cells
            lhdr[0].text = "שורה"
            lhdr[1].text = "טקסט"
            for cell in lhdr:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for r in p.runs:
                        r.bold = True

            for line_idx, line in enumerate(lines):
                row = ltable.add_row().cells
                line_b64 = line.get("image_b64") or ""
                if line_b64:
                    try:
                        line_bytes = base64.b64decode(line_b64)
                        # Clear default empty paragraph before inserting image.
                        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row[0].paragraphs[0].add_run().add_picture(
                            io.BytesIO(line_bytes), width=Inches(4.0)
                        )
                    except (ValueError, OSError) as e:
                        log.warning("line image decode failed: %s", e)
                        row[0].text = "[שגיאת תמונה]"
                line_text = str(line.get("text", ""))
                spans = spans_by_line.get(_line_key(page_num, line_idx), [])
                target_p = row[1].paragraphs[0]
                if spans:
                    _fill_paragraph_with_spans(target_p, line_text, spans)
                else:
                    target_p.add_run(line_text)
                for p in row[1].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "GET":
        return render_template("index.html")

    # Browser fallback for loading a saved JSON (multipart). The pywebview
    # build uses /load with a native file path instead.
    saved = request.files.get("saved")
    if saved and saved.filename:
        raw = saved.read()
        if not raw:
            return render_template("index.html", error="קובץ ה-JSON ריק")
        try:
            pages, filename, annotations, evaluation = parse_saved_json(raw)
        except (json.JSONDecodeError, ValueError, UnicodeDecodeError) as e:
            return render_template("index.html", error=f"שגיאה בקריאת JSON: {e}")
        return render_result(pages, filename, annotations, evaluation)

    return render_template("index.html", error="לא נבחר קובץ")


@app.route("/decode/start", methods=["POST"])
def decode_start():
    """Begin OCR. Accepts either a native file path (pywebview, JSON body)
    or a multipart 'pdf' upload (plain-browser fallback). Returns a job_id;
    progress is streamed from /decode/progress/<job_id>."""
    bad_type = "חייב להיות קובץ PDF או תמונה (jpg/png)"
    filename = "document.pdf"
    if request.is_json:
        if not DESKTOP_MODE:
            return jsonify(error="לא זמין בגרסת הווב"), 403
        body = request.get_json(silent=True) or {}
        path = body.get("path")
        if not path or not os.path.isfile(path):
            return jsonify(error="הקובץ לא נמצא"), 400
        if not path.lower().endswith(ACCEPTED_EXTS):
            return jsonify(error=bad_type), 400
        with open(path, "rb") as f:
            file_bytes = f.read()
        filename = os.path.basename(path)
        provider = body.get("model")
    else:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify(error="לא נבחר קובץ"), 400
        if not file.filename.lower().endswith(ACCEPTED_EXTS):
            return jsonify(error=bad_type), 400
        file_bytes = file.read()
        filename = file.filename
        provider = request.form.get("model")
    if not file_bytes:
        return jsonify(error="הקובץ ריק"), 400
    if provider not in MODELS:
        provider = DEFAULT_MODEL

    ext = os.path.splitext(filename)[1].lower()
    job_id = uuid.uuid4().hex[:12]
    JOBS[job_id] = {"q": queue.Queue(), "pages": None, "filename": filename}
    threading.Thread(
        target=run_job, args=(job_id, file_bytes, ext, provider), daemon=True
    ).start()
    return jsonify(job_id=job_id)


@app.route("/decode/progress/<job_id>")
def decode_progress(job_id):
    job = JOBS.get(job_id)
    if not job:
        return "unknown job", 404

    def stream():
        while True:
            ev = job["q"].get()
            yield f"data: {json.dumps(ev, ensure_ascii=False)}\n\n"
            if ev["type"] in ("end", "error"):
                break

    return Response(
        stream(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/result/<job_id>")
def result_page(job_id):
    job = JOBS.get(job_id)
    if not job or job.get("pages") is None:
        return redirect("/")
    return render_result(job["pages"], job["filename"])


@app.route("/load")
def load_saved():
    """Load a saved JSON by native path (pywebview only — reading an
    arbitrary server path is a file-read hole on a public server)."""
    if not DESKTOP_MODE:
        return render_template("index.html", error="לא זמין בגרסת הווב"), 403
    path = request.args.get("path")
    if not path or not os.path.isfile(path):
        return render_template("index.html", error="הקובץ לא נמצא")
    try:
        with open(path, "rb") as f:
            raw = f.read()
        pages, filename, annotations, evaluation = parse_saved_json(raw)
    except (json.JSONDecodeError, ValueError, UnicodeDecodeError, OSError) as e:
        return render_template("index.html", error=f"שגיאה בקריאת JSON: {e}")
    return render_result(pages, filename, annotations, evaluation)


@app.route("/rubrics", methods=["GET", "POST"])
def rubrics_endpoint():
    """GET → list. POST → create a new rubric ({name, content})."""
    if request.method == "GET":
        return jsonify(rubrics=list_rubrics())

    body = request.get_json(silent=True) or {}
    name = (body.get("name") or "").strip()
    content = (body.get("content") or "").strip()
    if not name:
        return jsonify(error="חסר שם רובריקה"), 400
    if not content:
        return jsonify(error="חסר תוכן רובריקה"), 400

    os.makedirs(RUBRICS_DIR, exist_ok=True)
    rubric_id = _slugify(name)
    path = os.path.join(RUBRICS_DIR, f"{rubric_id}.json")
    if os.path.exists(path):
        return jsonify(error="רובריקה בשם הזה כבר קיימת"), 409
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"name": name, "content": content}, f, ensure_ascii=False, indent=2)
    return jsonify(id=rubric_id, name=name)


@app.route("/rubrics/<rubric_id>")
def get_rubric(rubric_id):
    rubric = load_rubric(rubric_id)
    if not rubric:
        return jsonify(error="רובריקה לא נמצאה"), 404
    return jsonify(rubric)


@app.route("/evaluate", methods=["POST"])
def evaluate_endpoint():
    """Evaluate the current transcript against a rubric.

    Body: {pages: [...], rubric_id: "...", model: "claude"|"gemini"}.
    The browser already has the pages (from the rendered result), so it
    pushes them back — avoids parking them in JOBS for a separate fetch.
    """
    body = request.get_json(silent=True) or {}
    pages = body.get("pages") or []
    rubric_id = (body.get("rubric_id") or "").strip()
    provider = body.get("model") if body.get("model") in MODELS else DEFAULT_MODEL

    if not pages:
        return jsonify(error="אין טקסט לבדיקה"), 400
    rubric = load_rubric(rubric_id)
    if not rubric:
        return jsonify(error="רובריקה לא נמצאה"), 404

    try:
        result = evaluate_with_rubric(pages, rubric, provider)
    except Exception as e:
        log.exception("evaluate_endpoint failed")
        err = humanize_error(e, action="ההערכה")
        # Keep `error` key for old callers that just read .error; new fields
        # let the client render the richer banner.
        return jsonify(error=err["message"], **err), 502

    result["rubric_id"] = rubric_id
    result["rubric_name"] = rubric.get("name", rubric_id)
    highlights = _highlights_for_template(pages, result)
    return jsonify(evaluation=result, highlights=highlights)


# In-memory store for streaming eval jobs. Same lifecycle as the OCR JOBS
# dict — created on /evaluate/start, drained by the SSE handler, then dropped.
EVAL_JOBS: dict[str, dict] = {}


def run_evaluate_job(
    job_id: str, pages: list[dict], rubric: dict, provider: str, rubric_id: str
):
    job = EVAL_JOBS[job_id]
    q: queue.Queue = job["q"]
    try:
        for ev in evaluate_with_rubric_stream(pages, rubric, provider):
            if ev["type"] == "result":
                evaluation = ev["evaluation"]
                evaluation["rubric_id"] = rubric_id
                evaluation["rubric_name"] = rubric.get("name", rubric_id)
                highlights = _highlights_for_template(pages, evaluation)
                q.put({
                    "type": "done",
                    "evaluation": evaluation,
                    "highlights": highlights,
                })
            else:
                q.put(ev)
    except Exception as e:
        log.exception("run_evaluate_job failed")
        q.put({"type": "error", **humanize_error(e, action="ההערכה")})
    finally:
        q.put({"type": "end"})


@app.route("/evaluate/start", methods=["POST"])
def evaluate_start():
    """Streaming variant of /evaluate. Same body shape; returns a job_id and
    the caller subscribes to /evaluate/progress/<job_id> for per-criterion
    updates and the final result."""
    body = request.get_json(silent=True) or {}
    pages = body.get("pages") or []
    rubric_id = (body.get("rubric_id") or "").strip()
    provider = body.get("model") if body.get("model") in MODELS else DEFAULT_MODEL

    if not pages:
        return jsonify(error="אין טקסט לבדיקה"), 400
    rubric = load_rubric(rubric_id)
    if not rubric:
        return jsonify(error="רובריקה לא נמצאה"), 404

    job_id = uuid.uuid4().hex[:12]
    EVAL_JOBS[job_id] = {"q": queue.Queue()}
    threading.Thread(
        target=run_evaluate_job,
        args=(job_id, pages, rubric, provider, rubric_id),
        daemon=True,
    ).start()
    return jsonify(job_id=job_id)


@app.route("/evaluate/progress/<job_id>")
def evaluate_progress(job_id):
    job = EVAL_JOBS.get(job_id)
    if not job:
        return "unknown job", 404

    def stream():
        try:
            while True:
                ev = job["q"].get()
                yield f"data: {json.dumps(ev, ensure_ascii=False)}\n\n"
                if ev["type"] in ("end", "error"):
                    break
        finally:
            # One-shot: drop the job once the stream ends so memory doesn't
            # grow on repeated evaluations.
            EVAL_JOBS.pop(job_id, None)

    return Response(
        stream(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/evaluation/docx", methods=["POST"])
def evaluation_docx():
    """Browser fallback: streams a .docx of the evaluation back as a
    download. Desktop uses the native save_docx dialog instead and posts
    the same body shape, getting the same bytes."""
    body = request.get_json(silent=True) or {}
    evaluation = body.get("evaluation")
    filename = body.get("filename") or "תרגיל"
    rubric_name = body.get("rubric_name") or (evaluation or {}).get("rubric_name", "")
    pages = body.get("pages") or None
    if not evaluation:
        return jsonify(error="אין הערכה לייצוא"), 400
    try:
        data = build_evaluation_docx(evaluation, filename, rubric_name, pages=pages)
    except Exception as e:
        return jsonify(error=f"שגיאה ביצירת קובץ Word: {e}"), 500

    base = re.sub(r"\.(pdf|jpe?g|png|json)$", "", filename, flags=re.I)
    out_name = f"haskala-eval-{base}.docx"
    return send_file(
        io.BytesIO(data),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=out_name,
    )


if __name__ == "__main__":
    app.run(debug=True, port=5050, threaded=True)
