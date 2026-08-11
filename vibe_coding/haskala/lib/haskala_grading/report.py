"""report.py — the shared Word report builder.

Takes an evaluation dict (the shape core.EVAL_SCHEMA describes) plus the page
list, returns .docx bytes. Deliberately imports nothing from core, so the two
modules have no cycle and this one can be used on its own.

Lives in haskala/lib/ and is vendored into a product directory at deploy time —
see core.py's module docstring. The single source of truth is lib/.

Ported near-verbatim from haskala/products/checker/app.py's build_evaluation_docx
and its helpers (colored per-criterion highlighting, issue-span resolution,
rubric grid table) — see core.py's module docstring for what was deliberately
left out (decode-only export, numbered-overlay line crops).
"""
from __future__ import annotations

import base64
import difflib
import hashlib
import io
import logging
import re

log = logging.getLogger("haskala-grading")


# ─── languages / labels (mirrors checker/app.py DOCX_LABELS) ───────────────

DEFAULT_EXERCISE_LANG = "en"
# Mirrors core.DEFAULT_FEEDBACK_LANG: feedback follows the exercise language.
DEFAULT_FEEDBACK_LANG = DEFAULT_EXERCISE_LANG

LANGS: dict[str, dict] = {
    "he": {"name_he": "עברית",  "name_native": "עברית",   "dir": "rtl", "illegible": "[לא קריא]"},
    "en": {"name_he": "אנגלית", "name_native": "English", "dir": "ltr", "illegible": "[illegible]"},
    "ar": {"name_he": "ערבית",  "name_native": "العربية", "dir": "rtl", "illegible": "[غير مقروء]"},
}

DOCX_LABELS: dict[str, dict] = {
    "he": {
        "title": "בדיקת תרגיל — השכלה",
        "file": "קובץ:", "rubric": "רובריקה:", "overall_score": "ציון כללי:",
        "answered_yes": "התלמיד ענה על השאלה ✓",
        "answered_no": "התלמיד לא ענה על השאלה ✗",
        "per_criterion": "פירוט לפי קריטריון",
        "th_criterion": "קריטריון", "th_score": "ציון", "th_feedback": "פידבק",
        "feedback_secondary": "פידבק ({lang})",
        "word_count": "סך מילים בתרגיל: {n}",
        "helpful_words": "מילות עזר בשימוש",
        "helpful_words_line": "{count}/{total} מילות עזר נוצלו בתרגיל.",
        "used": "נוצלו: {words}", "not_used": "לא נוצלו: {words}",
        "summary": "סיכום",
        "original_document": "המסמך המקורי", "page": "עמוד {n}",
        "th_line": "שורה", "th_text": "טקסט", "image_error": "[שגיאת תמונה]",
        "text_analysis": "ניתוח טקסט",
        "wc_criterion": "ספירת מילים",
        "wc_detail": "נספרו {counted} מילים (נדרש {required}).",
        "wc_detail_no_rule": "נספרו {counted} מילים.",
    },
    "en": {
        "title": "Exercise Review — Haskala",
        "file": "File:", "rubric": "Rubric:", "overall_score": "Overall score:",
        "answered_yes": "The student answered the question ✓",
        "answered_no": "The student did not answer the question ✗",
        "per_criterion": "Per-criterion breakdown",
        "th_criterion": "Criterion", "th_score": "Score", "th_feedback": "Feedback",
        "feedback_secondary": "Feedback ({lang})",
        "word_count": "Total words: {n}",
        "helpful_words": "Helpful words used",
        "helpful_words_line": "{count}/{total} helpful words were used in the exercise.",
        "used": "Used: {words}", "not_used": "Not used: {words}",
        "summary": "Summary",
        "original_document": "Original document", "page": "Page {n}",
        "th_line": "Line", "th_text": "Text", "image_error": "[image error]",
        "text_analysis": "Text analysis",
        "wc_criterion": "Word count",
        "wc_detail": "{counted} words counted (required {required}).",
        "wc_detail_no_rule": "{counted} words counted.",
    },
    "ar": {
        "title": "مراجعة التمرين — هَسكالاه",
        "file": "الملف:", "rubric": "السلّم:", "overall_score": "التقييم العام:",
        "answered_yes": "أجاب الطالب عن السؤال ✓",
        "answered_no": "لم يُجب الطالب عن السؤال ✗",
        "per_criterion": "تفصيل حسب المعيار",
        "th_criterion": "المعيار", "th_score": "الدرجة", "th_feedback": "ملاحظات",
        "feedback_secondary": "ملاحظات ({lang})",
        "word_count": "عدد الكلمات: {n}",
        "helpful_words": "الكلمات المساعدة المستخدمة",
        "helpful_words_line": "{count}/{total} كلمات مساعدة استُخدمت في التمرين.",
        "used": "استُخدمت: {words}", "not_used": "لم تُستخدم: {words}",
        "summary": "ملخّص",
        "original_document": "المستند الأصلي", "page": "صفحة {n}",
        "th_line": "سطر", "th_text": "نص", "image_error": "[خطأ في الصورة]",
        "text_analysis": "تحليل النص",
        "wc_criterion": "عدد الكلمات",
        "wc_detail": "احتُسبت {counted} كلمة (المطلوب {required}).",
        "wc_detail_no_rule": "احتُسبت {counted} كلمة.",
    },
}

# Feature flag: the legacy compact per-criterion score/feedback table.
SHOW_LEGACY_EVAL_TABLE = False


# ─── criterion colors (mirrors checker/app.py:2005-2067) ───────────────────

_PALETTE: list[str] = [
    "#d9534f",  # 0 red
    "#2c7be5",  # 1 blue
    "#28a745",  # 2 green
    "#f0a000",  # 3 amber (richer than #e0a800 so it reads on cream paper)
    "#8e44ad",  # 4 purple
    "#e67e22",  # 5 orange
    "#16a085",  # 6 teal
    "#d63384",  # 7 magenta
    "#8B4513",  # 8 brown
    "#1a237e",  # 9 navy
    "#689f38",  # 10 lime
]

# Top-of-palette pins for the ministry rubric (checker's default). Order
# from most-frequent inline marks (mechanics: every spelling slip) to least
# (content: a few essay-level notes). Keyword lists tolerate the LLM
# returning slight renames (e.g. "Content & Organization", or a Hebrew title).
_PINNED_CRITERION_RULES: list[tuple[list[str], int]] = [
    # Slot 0 (red) — Mechanics / spelling / punctuation
    (["mechanics", "spelling", "punctuation",
      "איות", "כתיב", "פיסוק"], 0),
    # Slot 1 (blue) — Language Use / grammar / syntax
    (["language use", "grammar", "syntax",
      "תחביר", "דקדוק"], 1),
    # Slot 2 (green) — Vocabulary / lexical
    (["vocabulary", "lexical", "lexicon", "word choice",
      "אוצר מילים"], 2),
    # Slot 3 (amber) — Content / Organization / ideas / structure
    (["content", "organization", "structure", "ideas",
      "תוכן", "ארגון", "מבנה", "רעיונות"], 3),
]

# md5 fallback only picks from slots 4..end, so a non-pinned criterion can
# never accidentally land on a pinned color.
_FALLBACK_RANGE = list(range(4, len(_PALETTE)))


def color_for_criterion(name: str) -> str:
    """Hex color for a rubric criterion. Pinned slots first, else a
    deterministic md5 → palette slot from the remaining colors so the same
    criterion name always lands on the same color across runs."""
    n = (name or "").lower()
    for keywords, slot in _PINNED_CRITERION_RULES:
        if any(k.lower() in n for k in keywords):
            return _PALETTE[slot]
    digest = hashlib.md5(n.encode("utf-8")).hexdigest()
    slot = _FALLBACK_RANGE[int(digest, 16) % len(_FALLBACK_RANGE)]
    return _PALETTE[slot]


def attach_colors(evaluation: dict | None) -> dict | None:
    """Re-derive the 'color' field on every criterion from its name."""
    if not evaluation:
        return evaluation
    for c in evaluation.get("criteria") or []:
        c["color"] = color_for_criterion(c.get("name", ""))
    return evaluation


# Rubric-grid levels: four named bands from best to worst, matched against
# the evaluation's `level` field.
_GRID_LEVELS = [
    ("excellent", "Excellent"),
    ("good", "Good"),
    ("fair", "Fair"),
    ("needs_improvement", "Needs Improvement"),
]

# Normalization for matching a criteria_grid row to its evaluated criterion:
# lowercase, keep only ASCII alphanumerics + Hebrew.
_NORM_CRIT_RE = re.compile(r"[^a-z0-9֐-׿]+")


def _norm_crit(s: str) -> str:
    return _NORM_CRIT_RE.sub("", str(s or "").lower())


def _coerce_num(v):
    """Best-effort numeric coercion; None when the value isn't a number."""
    if isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return v
    try:
        return float(str(v).strip())
    except (TypeError, ValueError):
        return None


def _fmt_num(v) -> str:
    """Render a number without a trailing .0 (8.0 → '8', 7.5 → '7.5')."""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


# ─── issue span resolution (mirrors checker/app.py:2107-2191) ──────────────

_LINE_REF_RE = re.compile(r"p(\d+)-l(\d+)")

# Text the student should remove is shaded this yellow instead of its
# criterion's colour, per Avishai (2026-08-11). Softer than pure FFFF00, which
# is unreadable behind black text at report font sizes.
DELETE_HIGHLIGHT = "FFF176"
# Fuzzy fallback only forgives near-exact drift (punctuation, trailing space,
# 1-2 OCR slips). The model is instructed to return verbatim quotes, so loose
# matches that would paint the wrong span are rejected.
_FUZZY_MIN_BLOCK_RATIO = 0.8


def _find_span(line_text: str, quote: str) -> tuple[int, int] | None:
    """Locate `quote` within `line_text`. Tries exact, then case-insensitive,
    then a difflib longest-block fallback for small punctuation drift.
    Returns (start, end) on the original string, or None if no good match."""
    if not quote or not line_text:
        return None
    idx = line_text.find(quote)
    if idx >= 0:
        return idx, idx + len(quote)
    lower_line = line_text.lower()
    lower_quote = quote.lower()
    idx = lower_line.find(lower_quote)
    if idx >= 0:
        return idx, idx + len(lower_quote)
    matcher = difflib.SequenceMatcher(None, line_text, quote, autojunk=False)
    block = matcher.find_longest_match(0, len(line_text), 0, len(quote))
    min_size = max(4, int(len(quote) * _FUZZY_MIN_BLOCK_RATIO))
    if block.size < min_size:
        return None
    return block.a, block.a + block.size


def resolve_issue_spans(
    pages: list[dict], evaluation: dict | None
) -> dict[tuple[int, int], list[dict]]:
    """Map every (page, line_idx) to a list of highlight spans built from
    the evaluation's per-criterion issues. First criterion wins on overlap
    — a span that intersects an already-marked range is dropped."""
    if not evaluation:
        return {}
    spans: dict[tuple[int, int], list[dict]] = {}
    line_text_by_key: dict[tuple[int, int], str] = {}
    for page in pages:
        page_num = int(page.get("page", 0)) if str(page.get("page", "")).isdigit() else page.get("page", 0)
        for line_idx, line in enumerate(page.get("lines", [])):
            line_text_by_key[(int(page_num), line_idx)] = line.get("text", "")

    for crit in evaluation.get("criteria") or []:
        color = crit.get("color") or color_for_criterion(crit.get("name", ""))
        crit_name = crit.get("name", "")
        for issue in crit.get("issues") or []:
            ref = issue.get("line_ref", "")
            m = _LINE_REF_RE.search(ref)
            if not m:
                log.warning("issue with unparseable line_ref %r — skipped", ref)
                continue
            key = (int(m.group(1)), int(m.group(2)))
            line_text = line_text_by_key.get(key)
            if line_text is None:
                log.warning("issue references missing line %s — skipped", ref)
                continue
            quote = issue.get("quote", "")
            span = _find_span(line_text, quote)
            if not span:
                log.warning(
                    "issue quote not found in %s: %r — skipped", ref, quote[:50]
                )
                continue
            start, end = span
            existing = spans.setdefault(key, [])
            if any(not (end <= s["start"] or start >= s["end"]) for s in existing):
                continue  # first criterion wins on overlap
            existing.append(
                {
                    "start": start,
                    "end": end,
                    "color": color,
                    "comment": issue.get("comment", ""),
                    "criterion": crit_name,
                    "delete": bool(issue.get("delete")),
                }
            )

    for key in spans:
        spans[key].sort(key=lambda s: s["start"])
    return spans


# ─── DOCX export (mirrors checker/app.py:2196-2682) ─────────────────────────

def _tint_hex(hex_color: str, mix: float = 0.18) -> str:
    """Blend hex_color toward white. mix=0.18 → ~82% white tint, light
    enough to keep black text readable on top."""
    c = hex_color.lstrip("#")
    r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
    r = int(r * mix + 255 * (1 - mix))
    g = int(g * mix + 255 * (1 - mix))
    b = int(b * mix + 255 * (1 - mix))
    return f"{r:02x}{g:02x}{b:02x}"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Fill a table cell with a background color. python-docx exposes no
    helper for this, so we drop to the OOXML layer."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_run_shading(run, hex_color: str) -> None:
    """Background shade a single run (used for highlighting a substring
    inside a paragraph). python-docx has no API for run-level shading."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    rPr.append(shd)


def _fill_paragraph_with_spans(
    paragraph, text: str, spans: list[dict], include_notes: bool = True
) -> None:
    """Empty `paragraph`, then add runs reflecting `spans` as colored
    highlights over `text`. When `include_notes` is True, append an
    inline `(criterion: comment)` run shaded in the same color after
    every highlighted span. Spans are assumed sorted, non-overlapping."""
    from docx.shared import Pt

    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)
    if not text:
        return
    cursor = 0
    for span in spans:
        s, e = span["start"], span["end"]
        if s > cursor:
            paragraph.add_run(text[cursor:s])
        # A deletion is marked in one fixed way regardless of which criterion
        # raised it — brackets, a strikethrough and yellow — because it is an
        # instruction to the student rather than a category of mistake.
        # Avishai's rule (2026-08-11): yellow REPLACES the criterion colour.
        # Two shades over one span is unreadable, and the criterion is still
        # named in the note that follows.
        if span.get("delete"):
            tint = DELETE_HIGHLIGHT
            run = paragraph.add_run(f"[{text[s:e]}]")
            run.font.strike = True
        else:
            tint = _tint_hex(span["color"], mix=0.35)
            run = paragraph.add_run(text[s:e])
        _set_run_shading(run, tint)
        if include_notes:
            # The span's color already identifies the criterion, so the note
            # carries only the problem itself — Avishai's rule (2026-07-30).
            # The criterion name is kept as a fallback for the (rare) issue the
            # model returned with no comment at all.
            comment = span.get("comment", "") or span.get("criterion", "")
            if comment:
                note_run = paragraph.add_run(f" ({comment})")
                note_run.italic = True
                note_run.font.size = Pt(9)
                _set_run_shading(note_run, tint)
        cursor = e
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _word_count_cells(evaluation: dict, LBL: dict) -> tuple[str, str] | None:
    """The word-count row's (detail, marks) text, or None when the evaluation
    carries no word count at all.

    `length` is filled in by core.apply_length_deduction from the rubric's own
    word_count_rule; when a rubric declares no rule there is no deduction to
    show and the row degrades to a plain count.
    """
    length = evaluation.get("length") or {}
    counted = length.get("counted", evaluation.get("word_count"))
    if counted is None:
        return None
    required = length.get("required")
    detail = (
        LBL["wc_detail"].format(counted=counted, required=required)
        if required
        else LBL["wc_detail_no_rule"].format(counted=counted)
    )
    deduction = _coerce_num(length.get("deduction")) or 0
    marks = f"-{_fmt_num(deduction)}" if deduction else "0"
    return detail, marks


def _add_word_count_row(table, evaluation: dict, LBL: dict, ncols: int,
                        detail_col: int, marks_col: int, align) -> None:
    """Append the word-count row to whichever criterion table was rendered.
    Avishai's rule (2026-07-30): the length check is a row of the table like
    any other criterion, not a loose paragraph underneath it."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cells = _word_count_cells(evaluation, LBL)
    if cells is None:
        return
    detail, marks = cells
    row = table.add_row().cells
    row[0].text = LBL["wc_criterion"]
    for i in range(1, ncols):
        row[i].text = ""
    row[detail_col].text = detail
    row[marks_col].text = marks
    for cell in row:
        for p in cell.paragraphs:
            p.alignment = align
    _set_cell_shading(row[0], "e8e8e8")
    for p in row[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in row[marks_col].paragraphs:
        for r in p.runs:
            r.bold = True


def _add_text_analysis(doc, pages: list[dict], spans_by_line: dict,
                       LBL: dict, main_align) -> None:
    """The annotated exercise: a line-by-line table carrying the color
    highlights and the inline note on every marked span. Opens the report
    (Avishai, 2026-07-30) — it is what a teacher reads first; the criterion
    table below is the justification for the score.

    The original scan is deliberately NOT here. It goes at the end of the
    report (Ori, 2026-07-31): a full-width image between the heading and the
    analysis pushed the actual findings off the first screen."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading(LBL["text_analysis"], level=2).alignment = main_align

    for page in pages:
        page_num = page.get("page", "?")
        doc.add_heading(LBL["page"].format(n=page_num), level=3).alignment = main_align

        lines = page.get("lines") or []
        if not lines:
            continue
        # oris-scanner never carries per-line crops (English flow, no
        # numbered-box overlay) — always a single text column.
        ltable = doc.add_table(rows=1, cols=1)
        ltable.style = "Light Grid Accent 1"
        lhdr = ltable.rows[0].cells
        lhdr[0].text = LBL["th_text"]
        for cell in lhdr:
            for p in cell.paragraphs:
                p.alignment = main_align
                for r in p.runs:
                    r.bold = True

        for line_idx, line in enumerate(lines):
            row = ltable.add_row().cells
            line_text = str(line.get("text", ""))
            try:
                key = (int(page_num), line_idx)
            except (TypeError, ValueError):
                key = (page_num, line_idx)
            spans = spans_by_line.get(key, [])
            target_p = row[0].paragraphs[0]
            if spans:
                _fill_paragraph_with_spans(target_p, line_text, spans)
            else:
                target_p.add_run(line_text)
            for p in row[0].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_original_pages(doc, pages: list[dict], LBL: dict, main_align) -> None:
    """The student's own scan, one image per page, closing the report (Ori,
    2026-07-31). It is the evidence a teacher turns to when they doubt a
    finding, not something they read first — so it sits after the analysis and
    the criterion table rather than ahead of them."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    with_images = [p for p in pages if p.get("original_b64")]
    if not with_images:
        return

    doc.add_heading(LBL["original_document"], level=2).alignment = main_align
    for page in with_images:
        page_num = page.get("page", "?")
        doc.add_heading(LBL["page"].format(n=page_num), level=3).alignment = main_align
        try:
            orig_bytes = base64.b64decode(page["original_b64"])
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_para.add_run().add_picture(io.BytesIO(orig_bytes), width=Inches(5.5))
        except (ValueError, OSError) as e:
            log.warning("page %s original image failed: %s", page_num, e)


def build_evaluation_docx(
    evaluation: dict,
    filename: str,
    rubric_name: str,
    pages: list[dict] | None = None,
    feedback_lang: str = DEFAULT_FEEDBACK_LANG,
    exercise_lang: str = DEFAULT_EXERCISE_LANG,
) -> bytes:
    """Word document with the evaluation table + overall feedback. When
    `pages` is provided, appends each page's original scan followed by a
    line-by-line table (image | transcribed text, colored by issue) so the
    teacher can verify the OCR alongside the rubric scoring.

    Hebrew-friendly: paragraphs aligned right; the document is built top-down
    so existing readers (Google Docs, Word, LibreOffice) all open it cleanly.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    LBL = DOCX_LABELS.get(feedback_lang, DOCX_LABELS["he"])
    fb = LANGS.get(feedback_lang, LANGS[DEFAULT_FEEDBACK_LANG])
    main_align = (
        WD_ALIGN_PARAGRAPH.LEFT if fb["dir"] == "ltr" else WD_ALIGN_PARAGRAPH.RIGHT
    )

    doc = Document()

    title = doc.add_heading(LBL["title"], level=1)
    title.alignment = main_align

    meta = doc.add_paragraph()
    meta.alignment = main_align
    meta.add_run(f"{LBL['file']} {filename}\n").bold = True
    meta.add_run(f"{LBL['rubric']} {rubric_name}\n").bold = True
    meta.add_run(f"{LBL['overall_score']} {evaluation.get('overall_score', '')}").bold = True

    qa = str(evaluation.get("question_answered", "")).strip().lower()
    if qa in ("yes", "no"):
        ans_p = doc.add_paragraph()
        ans_p.alignment = main_align
        ans_p.add_run(LBL["answered_yes"] if qa == "yes" else LBL["answered_no"]).bold = True

    criteria = attach_colors(evaluation).get("criteria") or []

    # Text analysis first — the annotated exercise, then the scoring tables.
    spans_by_line = resolve_issue_spans(pages, evaluation) if pages else {}
    if pages:
        _add_text_analysis(doc, pages, spans_by_line, LBL, main_align)

    show_secondary = feedback_lang != exercise_lang
    secondary_lang = LANGS.get(exercise_lang, LANGS[DEFAULT_EXERCISE_LANG])
    secondary_header = LBL["feedback_secondary"].format(lang=secondary_lang["name_native"])
    secondary_align_ltr = secondary_lang["dir"] == "ltr"

    def _secondary_text(c: dict) -> str:
        return str(c.get("feedback_secondary") or c.get("feedback_en") or "")

    # Rubric grid (criteria × 4 levels + marks). Only rendered when the
    # rubric carried a structured `criteria_grid`.
    grid = evaluation.get("criteria_grid")
    has_grid = isinstance(grid, list) and len(grid) > 0
    if has_grid:
        gcols = 1 + len(_GRID_LEVELS) + 1  # criterion | levels… | marks
        gtable = doc.add_table(rows=1, cols=gcols)
        gtable.style = "Light Grid Accent 1"
        ghdr = gtable.rows[0].cells
        ghdr[0].text = "Criteria"
        for i, (_key, label) in enumerate(_GRID_LEVELS):
            ghdr[1 + i].text = label
        ghdr[gcols - 1].text = "Marks"
        for cell in ghdr:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True

        for gi, g in enumerate(grid):
            gname = _norm_crit(g.get("name"))
            ev = next((c for c in criteria if _norm_crit(c.get("name")) == gname), None)
            if ev is None and gi < len(criteria):
                ev = criteria[gi]
            color = (ev and ev.get("color")) or color_for_criterion(g.get("name", ""))
            max_marks = _coerce_num(g.get("max_marks"))
            if max_marks is None and ev is not None:
                max_marks = _coerce_num(ev.get("max_score"))
            score = _coerce_num(ev.get("score")) if ev else None

            level_idx = -1
            if ev:
                lvl = str(ev.get("level") or "").lower()
                level_idx = next(
                    (i for i, (k, _l) in enumerate(_GRID_LEVELS) if k == lvl), -1
                )
            if level_idx < 0 and max_marks and score is not None:
                level_idx = round((1 - score / max_marks) / 0.25)
                level_idx = min(3, max(0, level_idx))

            row = gtable.add_row().cells
            row[0].text = str(g.get("name") or "")
            _set_cell_shading(row[0], _tint_hex(color))
            for p in row[0].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True

            levels = g.get("levels") or {}
            for i, (key, _label) in enumerate(_GRID_LEVELS):
                cell = row[1 + i]
                cell.text = str(levels.get(key) or "")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                if i == level_idx:
                    _set_cell_shading(cell, _tint_hex(color, 0.35))

            marks_cell = row[gcols - 1]
            score_txt = "" if score is None else _fmt_num(score)
            max_txt = "" if max_marks is None else _fmt_num(max_marks)
            marks_cell.text = f"{score_txt}/{max_txt}"
            marks_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        _add_word_count_row(gtable, evaluation, LBL, gcols,
                            detail_col=1, marks_col=gcols - 1,
                            align=WD_ALIGN_PARAGRAPH.LEFT)

    # Legacy per-criterion score/feedback table. Shown when no grid is
    # available, or forced on via SHOW_LEGACY_EVAL_TABLE even when a grid
    # exists.
    if SHOW_LEGACY_EVAL_TABLE or not has_grid:
        doc.add_heading(LBL["per_criterion"], level=2).alignment = main_align
        cols = 4 if show_secondary else 3
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text = LBL["th_criterion"]
        header[1].text = LBL["th_score"]
        header[2].text = LBL["th_feedback"]
        if show_secondary:
            header[3].text = secondary_header
        for idx, cell in enumerate(header):
            align = (
                WD_ALIGN_PARAGRAPH.LEFT
                if idx == 3 and secondary_align_ltr
                else main_align
            )
            for p in cell.paragraphs:
                p.alignment = align
                for r in p.runs:
                    r.bold = True

        for c in criteria:
            row = table.add_row().cells
            row[0].text = str(c.get("name", ""))
            row[1].text = f"{c.get('score', '')}/{c.get('max_score', '')}"
            row[2].text = str(c.get("feedback", ""))
            if show_secondary:
                row[3].text = _secondary_text(c)
            for idx, cell in enumerate(row):
                align = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if idx == 3 and secondary_align_ltr
                    else WD_ALIGN_PARAGRAPH.RIGHT
                )
                for p in cell.paragraphs:
                    p.alignment = align
            color = c.get("color") or color_for_criterion(c.get("name", ""))
            _set_cell_shading(row[0], _tint_hex(color))
            for p in row[0].paragraphs:
                for r in p.runs:
                    r.bold = True

        # Bottom row of the criterion table: words counted under the rubric's
        # rules and the points the length cost. Replaces the loose word-count
        # paragraph that used to sit below the table.
        _add_word_count_row(table, evaluation, LBL, cols,
                            detail_col=2, marks_col=1, align=main_align)

    hw = evaluation.get("helpful_words_usage")
    if hw and hw.get("total"):
        doc.add_heading(LBL["helpful_words"], level=2).alignment = main_align
        hw_p = doc.add_paragraph()
        hw_p.alignment = main_align
        hw_p.add_run(LBL["helpful_words_line"].format(
            count=hw.get("count", 0), total=hw.get("total", 0))).bold = True
        if hw.get("used"):
            used_p = doc.add_paragraph(LBL["used"].format(words=", ".join(hw["used"])))
            used_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if hw.get("unused"):
            unused_p = doc.add_paragraph(LBL["not_used"].format(words=", ".join(hw["unused"])))
            unused_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # The prose summary ("סיכום" / overall_feedback) is deliberately not
    # rendered — Avishai's rule (2026-07-30). The model still produces it and
    # it is still logged, so it stays available for debugging a surprising
    # score without occupying space in the teacher-facing report.

    # The "clean transcript" and "marked transcript" sections that used to
    # follow are gone (Ori, 2026-07-31). Between them and the text analysis at
    # the top, the report carried the same transcript three times over; the
    # analysis is the one that carries the findings, so it is the one that
    # stayed.

    if pages:
        doc.add_page_break()
        # Last of all: the student's own scan.
        _add_original_pages(doc, pages, LBL, main_align)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
