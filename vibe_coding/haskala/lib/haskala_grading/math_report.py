"""math_report.py — בוני הדוחות של בדיקת המתמטיקה (HTML עצמאי + Word).

מקבל את רשימת ה-pages ש-`math_core.process_stream`/`check_file` מחזירים ומייצר
ממנה דוח. אין כאן קריאות AI ואין Flask, ואין import מ-math_core — כך ש-CLI או
poller יכולים לייצר דוח בלי לגרור את ה-pipeline, ואין מעגל import.

שני צרכנים:
  • products/math-checker  — ה-UI מציג HTML, ומייצא Word
  • products/math-form-checker — כותב Word לתיקיית ה-Drive של המורה

`approved` הוא ההבדל שחשוב ל-form-checker: הזרימה מהטופס היא חד-פעמית ואין בה
מסך שבו המורה מאשר ניקוד, ולכן היא תמיד מייצרת דוח עם approved=False — הכותרת
אומרת במפורש "טיוטה — טרם אושר ע\"י המורה", והניקוד מוצג כהצעה.
"""
from __future__ import annotations

import base64
import io
import logging
import re

log = logging.getLogger("haskala.math")

# ─── export builders (HTML + Word) ──────────────────────────────────────────

VERDICT_HE = {"correct": "נכון ✓", "partial": "חלקי", "incorrect": "שגוי ✗", "unclear": "לא ברור"}
VERDICT_COLOR = {"correct": "1e7e34", "partial": "a0740a", "incorrect": "b54343", "unclear": "5a6b82"}


def _num(v):
    """Coerce a points value to a number, or None if missing/blank/non-numeric."""
    if v is None or v == "":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def _fmt_pts(v) -> str:
    """Render a points number without a trailing .0 (7.0 → "7", 7.5 → "7.5")."""
    n = _num(v)
    if n is None:
        return "—"
    return str(int(n)) if n == int(n) else str(n)


def compute_totals(pages: list[dict]) -> tuple[float, float]:
    """Sum (earned, max) points across every problem on every page."""
    earned = total = 0.0
    for p in pages or []:
        for pr in ((p.get("analysis") or {}).get("problems") or []):
            mx = _num(pr.get("points_max"))
            if mx is None:
                continue
            total += mx
            earned += _num(pr.get("points_earned")) or 0
    return earned, total


# The per-item default in math_core's prompt when a page has no printed points
# and no rubric was supplied. Kept in sync by `test_default_points_max_matches`.
DEFAULT_POINTS_MAX = 10.0

# Every report states the overall grade on this scale, whatever the raw sum was.
GRADE_SCALE = 100.0


def compute_grade(pages: list[dict]) -> tuple[float | None, bool]:
    """The overall grade rescaled to GRADE_SCALE, and whether we chose the
    denominator ourselves.

    The raw sum is not a grade. Pinning the grading unit to the sub-part turned
    a six-page algebra exam into 24 items worth 240 points without a single
    judgement changing — and a teacher who set that exam out of 100 has no way
    to read "134/240". Rescaling puts the number back on the scale teachers
    actually use.

    The second value is True when every points_max is the prompt's 10-point
    default, i.e. no rubric was supplied and no points were printed on the page.
    The grade is then a percentage of a scale we invented, and the report says
    so rather than presenting our convention as the teacher's."""
    earned, total = compute_totals(pages)
    if total <= 0:
        return None, False
    maxes = [m for m in (_num(pr.get("points_max"))
                         for p in pages or []
                         for pr in ((p.get("analysis") or {}).get("problems") or []))
             if m is not None]
    ours = bool(maxes) and all(m == DEFAULT_POINTS_MAX for m in maxes)
    # Half points, because that is the finest granularity the model may award.
    return round(earned / total * GRADE_SCALE * 2) / 2, ours


def grade_basis_note(pages: list[dict], ours: bool) -> str:
    """One line under the grade saying what it was computed from.

    Without it the grade reads as the teacher's own scale in every case. The
    item count is part of the note because it is the number a teacher can check
    against the exam in front of them — if we split 12 exercises into 24, that
    is visible here and not only in a grade that looks slightly off."""
    n = sum(len((p.get("analysis") or {}).get("problems") or []) for p in pages or [])
    _, total = compute_totals(pages)
    if ours:
        return (f"מבוסס על {n} סעיפים בניקוד אחיד — לא סופק מחוון ולא נמצא ניקוד "
                f"מודפס בדף, ולכן משקל הסעיפים נקבע על ידינו.")
    return f"מבוסס על {n} סעיפים, {_fmt_pts(total)} נקודות לפי המחוון/הניקוד שבדף."


def build_result_html(pages: list[dict], filename: str, approved: bool = False) -> str:
    """Standalone, self-contained HTML report. KaTeX is pulled from a CDN and
    auto-renders \\(...\\) so the LaTeX shows as real math; the scan images are
    embedded as base64 so the file opens anywhere with no server. Mirrors the
    web UI layout (verdict badges, ok/bad step coloring, comments, feedback)."""
    import html as _html

    def esc(s) -> str:
        return _html.escape(str(s if s is not None else ""))

    def math(latex) -> str:
        # KaTeX auto-render reads textContent, so escaping is safe (and needed
        # for inequalities like a<b). Wrapped in \( \) inline delimiters.
        return r"\(" + esc(latex) + r"\)"

    body: list[str] = []
    for p in pages or []:
        a = p.get("analysis") or {}
        rot = p.get("rotation_applied") or 0
        rot_lbl = f" · סובב {rot}°" if rot else ""
        body.append(f'<section class="page"><h2>עמוד {esc(p.get("page", "?"))}'
                    f'<span class="rot">{esc(rot_lbl)}</span></h2><div class="layout">')
        if p.get("image_b64"):
            body.append(f'<div class="imgwrap"><img src="data:image/jpeg;base64,'
                        f'{p["image_b64"]}" alt="עמוד {esc(p.get("page", ""))}"></div>')
        body.append('<div class="analysis">')
        if a.get("page_summary"):
            body.append(f'<div class="summary">{esc(a["page_summary"])}</div>')
        if a.get("has_diagram"):
            body.append(f'<div class="diagram">▣ סרטוט: {esc(a.get("diagram_description", ""))}</div>')
        problems = a.get("problems") or []
        if not problems:
            body.append('<div class="noprob">לא זוהו תרגילים בעמוד זה.</div>')
        for pr in problems:
            v = pr.get("verdict", "unclear")
            body.append('<div class="prob"><div class="prob-head">'
                        f'<span class="prob-id">תרגיל {esc(pr.get("id", "?"))}</span>'
                        f'<span class="topic">{esc(pr.get("topic", ""))}</span>'
                        f'<span class="badge b-{esc(v)}">{esc(VERDICT_HE.get(v, v))}</span></div>')
            if pr.get("statement_latex"):
                body.append(f'<div class="step statement" dir="ltr">{math(pr["statement_latex"])}</div>')
            for s in pr.get("student_steps") or []:
                cls = "bad" if s.get("ok") is False else ("ok" if s.get("ok") is True else "")
                body.append(f'<div class="step {cls}" dir="ltr">{math(s.get("latex", ""))}')
                if s.get("comment"):
                    body.append(f'<div class="step-comment">⚠ {esc(s["comment"])}</div>')
                body.append('</div>')
            if pr.get("final_answer_latex"):
                body.append(f'<div class="step final" dir="ltr"><b>תשובה:</b> {math(pr["final_answer_latex"])}</div>')
            if _num(pr.get("points_max")) is not None:
                body.append(
                    '<div class="score-row"><span class="pts">ניקוד: '
                    f'{_fmt_pts(pr.get("points_earned"))} / {_fmt_pts(pr.get("points_max"))}'
                    '</span>'
                    + (f'<span class="score-note">{esc(pr["score_suggestion"])}</span>'
                       if pr.get("score_suggestion") else "")
                    + '</div>')
            elif pr.get("score_suggestion"):
                body.append(f'<div class="score-row">{esc(pr["score_suggestion"])}</div>')
            if pr.get("feedback"):
                body.append(f'<div class="feedback">{esc(pr["feedback"])}</div>')
            body.append('</div>')
        body.append('</div></div></section>')

    head = """<!DOCTYPE html>
<html dir="rtl" lang="he"><head><meta charset="utf-8">
<title>בדיקת מתמטיקה — __TITLE__</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"
  onload="renderMathInElement(document.body,{delimiters:[{left:'\\\\(',right:'\\\\)',display:false},{left:'\\\\[',right:'\\\\]',display:true}],throwOnError:false});"></script>
<style>
  body{font-family:'Segoe UI',Arial,sans-serif;color:#2b3442;background:#f4f1ea;margin:0;padding:2rem;line-height:1.5}
  .doc-title{font-size:1.5rem;font-weight:800;color:#2e7286;margin:0 0 0.3rem}
  .doc-meta{color:#6e6a5c;margin:0 0 1.5rem;font-size:0.9rem}
  .page{background:#fff;border:1px solid #d8cfb6;border-radius:12px;padding:1.2rem 1.4rem;margin-bottom:1.5rem;box-shadow:0 1px 4px rgba(0,0,0,.06)}
  .page h2{margin:0 0 1rem;color:#2e7286;font-size:1.2rem}
  .rot{color:#6e6a5c;font-size:0.85rem;font-weight:400}
  .layout{display:flex;gap:1.4rem;align-items:flex-start;flex-wrap:wrap}
  .imgwrap{flex:0 0 320px;max-width:340px}
  .imgwrap img{width:100%;border:1px solid #d8cfb6;border-radius:8px}
  .analysis{flex:1;min-width:300px}
  .summary{font-size:0.95rem;color:#3a4250;margin-bottom:0.6rem}
  .diagram{font-size:0.85rem;color:#2e7286;background:#eef4f7;border-radius:6px;padding:0.3rem 0.6rem;margin-bottom:0.6rem;display:inline-block}
  .noprob{color:#6e6a5c;font-style:italic}
  .prob{border:1px solid #e7e0cd;border-radius:10px;padding:0.8rem 1rem;margin-bottom:0.9rem}
  .prob-head{display:flex;gap:0.6rem;align-items:center;margin-bottom:0.5rem;flex-wrap:wrap}
  .prob-id{font-weight:700}
  .topic{font-size:0.78rem;color:#6e6a5c;background:#ece8db;border-radius:999px;padding:0.15rem 0.6rem}
  .badge{font-size:0.82rem;padding:0.18rem 0.7rem;border-radius:999px;font-weight:700}
  .b-correct{background:rgba(30,126,52,.15);color:#1e7e34}
  .b-partial{background:rgba(160,116,10,.15);color:#a0740a}
  .b-incorrect{background:rgba(181,67,67,.15);color:#b54343}
  .b-unclear{background:rgba(90,107,130,.15);color:#5a6b82}
  .step{padding:0.5rem 0.8rem;margin:0.3rem 0;border-radius:6px;border-inline-start:3px solid #d8cfb6;background:#f6f3eb;direction:ltr;text-align:left}
  .step.ok{border-inline-start-color:#1e7e34}
  .step.bad{border-inline-start-color:#b54343;background:#fff0f0}
  .step.statement{border-inline-start-color:#2e7286;background:#eef4f7}
  .step.final{border-inline-start-color:#b3924a;background:#faf5e9;font-weight:600}
  .step-comment{color:#b54343;font-size:0.85rem;margin-top:0.3rem;direction:rtl;text-align:right;font-weight:500}
  .score-row{color:#3a4250;font-size:0.9rem;margin:0.6rem 0 0.3rem;display:flex;gap:0.6rem;align-items:baseline;flex-wrap:wrap}
  .score-row .pts{font-weight:700;color:#2e7286}
  .score-row .score-note{color:#6e6a5c;font-size:0.85rem}
  .feedback{margin-top:0.5rem;padding-top:0.6rem;border-top:1px dashed #d8cfb6;font-size:0.93rem}
  .total-box{background:#fff;border:1px solid #d8cfb6;border-radius:12px;padding:1rem 1.4rem;margin-bottom:1.5rem;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:0.8rem}
  .total-grade{font-size:1.3rem;font-weight:800;color:#2e7286}
  .total-basis{display:block;margin:-1rem 0 1.5rem;font-size:0.85rem;color:#6b6459}
  .approve-badge{font-size:0.9rem;font-weight:700;padding:0.3rem 0.9rem;border-radius:999px}
  .approve-yes{background:rgba(30,126,52,.15);color:#1e7e34}
  .approve-no{background:rgba(160,116,10,.15);color:#a0740a}
  @media print{body{background:#fff;padding:0}.page{box-shadow:none;break-inside:avoid}}
</style></head><body>
<div class="doc-title">בדיקת מתמטיקה — השכלה</div>
<div class="doc-meta">קובץ: __TITLE__ · __NPAGES__ עמודים</div>
__TOTAL_BOX__
"""
    npages = len(pages or [])
    grade, ours = compute_grade(pages)
    if grade is not None:
        badge = ('<span class="approve-badge approve-yes">✓ אושר ע"י המורה</span>'
                 if approved else
                 '<span class="approve-badge approve-no">טיוטה — טרם אושר</span>')
        basis = (f'<span class="total-basis">{_html.escape(grade_basis_note(pages, ours))}</span>')
        total_box = (
            '<div class="total-box">'
            f'<span class="total-grade">ציון כולל: {_fmt_pts(grade)} / {_fmt_pts(GRADE_SCALE)}</span>'
            f'{badge}</div>{basis}')
    else:
        total_box = ""
    head = (head.replace("__TITLE__", _html.escape(filename))
                .replace("__NPAGES__", str(npages))
                .replace("__TOTAL_BOX__", total_box))
    return head + "\n".join(body) + "\n</body></html>"


# ─── LaTeX → Unicode, for the Word report ────────────────────────────────────
# The docx used to print the model's LaTeX verbatim, on the stated assumption
# that "the strings are simple, e.g. x^2-6x+10". That holds for algebra and
# collapses on geometry, which is most of what Avishai's classes actually
# submit: a real report line came out as
#
#     \triangle A_1=\triangle A_2,\ \triangle BLA \cong \triangle BLA
#
# which a teacher cannot read. The HTML report never had the problem because it
# renders with KaTeX; only the Word path is affected.
#
# Unicode rather than images or Word's own equation format (OMML), for three
# reasons that all point the same way:
#   • The maths in these reports is symbol-heavy but structurally flat —
#     congruences, angles, subscripts — which Unicode represents exactly.
#   • The text stays selectable and editable, which is the point of shipping
#     .docx at all: the teacher corrects the report before passing it on.
#   • LaTeX→OMML needs an XSLT that ships with Microsoft Office and is not
#     redistributable, and an image per step would bloat the file and lose the
#     ability to edit. Neither is worth it for △ and ≅.

_LATEX_SYMBOLS = {
    r"\triangle": "△", r"\angle": "∠", r"\measuredangle": "∡",
    r"\cong": "≅", r"\sim": "∼", r"\simeq": "≃", r"\equiv": "≡",
    r"\perp": "⊥", r"\parallel": "∥", r"\nparallel": "∦",
    r"\cdot": "·", r"\times": "×", r"\div": "÷", r"\pm": "±", r"\mp": "∓",
    r"\neq": "≠", r"\ne": "≠", r"\leq": "≤", r"\le": "≤",
    r"\geq": "≥", r"\ge": "≥", r"\approx": "≈", r"\propto": "∝",
    r"\infty": "∞", r"\degree": "°", r"\circ": "°",
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\theta": "θ", r"\lambda": "λ", r"\mu": "μ", r"\pi": "π",
    r"\sigma": "σ", r"\phi": "φ", r"\omega": "ω",
    r"\Delta": "Δ", r"\Sigma": "Σ", r"\Omega": "Ω",
    r"\rightarrow": "→", r"\to": "→", r"\leftarrow": "←",
    r"\Rightarrow": "⇒", r"\Leftrightarrow": "⇔", r"\implies": "⇒",
    r"\in": "∈", r"\notin": "∉", r"\subset": "⊂", r"\subseteq": "⊆",
    r"\cup": "∪", r"\cap": "∩", r"\emptyset": "∅", r"\varnothing": "∅",
    r"\forall": "∀", r"\exists": "∃", r"\therefore": "∴", r"\because": "∵",
    r"\ldots": "…", r"\dots": "…", r"\cdots": "⋯",
    r"\overline": "", r"\bar": "", r"\vec": "",
    r"\sum": "Σ", r"\prod": "∏", r"\int": "∫",
    r"\%": "%", r"\$": "$", r"\&": "&", r"\#": "#",
}

# Spacing and grouping commands that carry no glyph. Order matters: the longer
# names must be tried before their own prefixes ("\quad" before "\q...").
_LATEX_BLANKS = [r"\qquad", r"\quad", r"\left", r"\right", r"\displaystyle",
                 r"\limits", r"\,", r"\;", r"\:", r"\!", r"\ "]

_SUP = str.maketrans("0123456789+-=()aeionx", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ᵃᵉⁱᵒⁿˣ")
_SUB = str.maketrans("0123456789+-=()aeiox", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑᵢₒₓ")


def _brace_arg(s: str, i: int) -> tuple[str, int]:
    """The {...} group starting at s[i], and the index just past it.

    Falls back to the single character at s[i] when there is no brace, which is
    how LaTeX itself reads x^2 — and the model writes both forms."""
    if i >= len(s):
        return "", i
    if s[i] == "\\":
        # A command as the argument, e.g. 90^\circ or x^\alpha. Without this the
        # backslash alone is taken as the whole argument and the command name
        # spills into the surrounding text — a real report line came out as
        # "∠ALM = 90^circ".
        j = i + 1
        while j < len(s) and s[j].isalpha():
            j += 1
        return s[i:j], j
    if s[i] != "{":
        return s[i], i + 1
    depth, j = 0, i
    while j < len(s):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1:j], j + 1
        j += 1
    return s[i + 1:], len(s)      # unbalanced — take the rest rather than fail


def _script(body: str, table, fallback: str) -> str:
    """A super/subscript as Unicode when every character maps, else a readable
    ASCII form. Half-translated scripts are worse than none: 'x²ʸ' next to
    'x^(2y)' in the same report reads as two different expressions."""
    inner = latex_to_unicode(body)
    # "90^\circ" is how LaTeX spells 90°, and the degree sign is already raised —
    # emitting "90^°" would be both wrong and uglier than the source.
    if inner == "°":
        return "°"
    if inner and all(c in table for c in map(ord, inner)):
        return inner.translate(table)
    return f"{fallback}({inner})" if len(inner) > 1 else f"{fallback}{inner}"


def latex_to_unicode(latex: str) -> str:
    """Readable plain text for a LaTeX fragment produced by the analysis model.

    Not a general LaTeX engine and not trying to be — it covers what
    ANALYSIS_PROMPT actually asks for ("תוכן בלבד, ללא $"): relations, geometry
    symbols, scripts, fractions and roots. Anything unrecognised degrades to its
    own name without the backslash, which is still strictly more readable than
    the raw source and never raises.
    """
    if not latex:
        return ""
    s = str(latex)
    out: list[str] = []
    i = 0
    while i < len(s):
        ch = s[i]

        if ch == "\\":
            # \text{...} — the model uses it for Hebrew words inside formulas,
            # so the content passes through untouched rather than being parsed
            # as maths.
            for cmd in (r"\text", r"\mathrm", r"\mbox", r"\textrm"):
                if s.startswith(cmd + "{", i):
                    body, i = _brace_arg(s, i + len(cmd))
                    out.append(body)
                    break
            else:
                # \tfrac and \cfrac are here because the model does use them and
                # they are not cosmetic variants to us: an unhandled one falls
                # through to the strip-the-backslash branch, so a real report
                # line read "BC = tfrac12 AN" instead of "BC = 1/2 AN"
                # (production, 6/8/2026). All four render identically in plain
                # text — the display/inline distinction has no meaning here.
                frac = next((c for c in (r"\dfrac", r"\tfrac", r"\cfrac", r"\frac")
                             if s.startswith(c, i)), None)
                if frac:
                    num, i = _brace_arg(s, i + len(frac))
                    den, i = _brace_arg(s, i)
                    num_u, den_u = latex_to_unicode(num), latex_to_unicode(den)
                    num_s = num_u if len(num_u) <= 1 else f"({num_u})"
                    den_s = den_u if len(den_u) <= 1 else f"({den_u})"
                    out.append(f"{num_s}/{den_s}")
                elif s.startswith(r"\sqrt", i):
                    body, i = _brace_arg(s, i + 5)
                    body_u = latex_to_unicode(body)
                    out.append("√" + (body_u if len(body_u) <= 1 else f"({body_u})"))
                else:
                    for blank in _LATEX_BLANKS:
                        if s.startswith(blank, i):
                            out.append(" ")
                            i += len(blank)
                            break
                    else:
                        # Longest symbol name first, so \le does not shadow \leq.
                        for name in sorted(_LATEX_SYMBOLS, key=len, reverse=True):
                            if s.startswith(name, i):
                                out.append(_LATEX_SYMBOLS[name])
                                i += len(name)
                                # LaTeX ends a command name at the first space,
                                # and that space is a delimiter rather than
                                # something to print — without dropping it,
                                # "\triangle ABC" comes out as "△ ABC".
                                if i < len(s) and s[i] == " ":
                                    i += 1
                                break
                        else:
                            j = i + 1
                            while j < len(s) and s[j].isalpha():
                                j += 1
                            out.append(s[i + 1:j])   # unknown command, minus the \
                            i = j if j > i + 1 else i + 1
                            if i < len(s) and s[i] == " ":
                                i += 1
            continue

        if ch == "^":
            body, i = _brace_arg(s, i + 1)
            out.append(_script(body, _SUP, "^"))
            continue
        if ch == "_":
            body, i = _brace_arg(s, i + 1)
            out.append(_script(body, _SUB, "_"))
            continue
        if ch in "{}":
            i += 1
            continue

        out.append(ch)
        i += 1

    text = " ".join("".join(out).split())

    # Dropping the command-terminating space above is right for prefixes (△ABC)
    # and wrong for relations, which would come out as "△ABC ≅△DEF" — spaced on
    # one side only. Restoring a single space around the relation glyphs, and
    # only those, gives "△ABC ≅ △DEF". "=" is deliberately untouched: it is
    # typed literally, so whatever spacing the model chose is intentional.
    for glyph in "≅∥⊥≠≤≥≈≡∼→⇒⇔∈∉⊂⊆∪∩±∓×÷∝":
        text = re.sub(rf"\s*{re.escape(glyph)}\s*", f" {glyph} ", text)
    return " ".join(text.split())


def build_result_docx(pages: list[dict], filename: str, approved: bool = False) -> bytes:
    """Word (.docx) report. Math is converted from the model's LaTeX to Unicode
    by latex_to_unicode — see the note above it for why not images or OMML —
    and set LTR so it flows correctly inside the RTL document. Each page carries
    its scan image, verdicts, comments, score and feedback. Hebrew paragraphs
    are right-aligned + bidi for correct RTL flow."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def rtl(par):
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = par._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
        return par

    def math_run(par, latex):
        """One converted formula, marked explicitly as left-to-right.

        The w:rtl=0 is not decoration. Word inherits paragraph direction into
        runs, and inside an RTL document a bare formula gets reordered — the
        Drive importer doing exactly this is why these reports are .docx and not
        Google Docs in the first place. Marking the run LTR pins it.

        Cambria Math rather than Consolas: it carries △ ∠ ≅ ⊥ ∥ and the
        sub/superscript digits, which a code font renders as boxes on many
        machines — a report full of tofu is worse than the raw LaTeX was.
        """
        run = par.add_run(latex_to_unicode(latex))
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)
        rPr = run._r.get_or_add_rPr()
        rtl_el = OxmlElement("w:rtl")
        rtl_el.set(qn("w:val"), "0")
        rPr.append(rtl_el)
        return run

    doc = Document()
    rtl(doc.add_heading("בדיקת מתמטיקה — השכלה", level=1))
    meta = rtl(doc.add_paragraph())
    meta.add_run(f"קובץ: {filename}").bold = True

    grade, ours = compute_grade(pages)
    if grade is not None:
        tp = rtl(doc.add_paragraph())
        tr = tp.add_run(f"ציון כולל: {_fmt_pts(grade)} / {_fmt_pts(GRADE_SCALE)}")
        tr.bold = True
        tr.font.size = Pt(14)
        tr.font.color.rgb = RGBColor(0x2E, 0x72, 0x86)
        bp = rtl(doc.add_paragraph())
        br = bp.add_run(grade_basis_note(pages, ours))
        br.font.size = Pt(9)
        br.font.color.rgb = RGBColor(0x6B, 0x64, 0x59)
        sp = rtl(doc.add_paragraph())
        sr = sp.add_run('✓ אושר ע"י המורה' if approved else "טיוטה — טרם אושר ע\"י המורה")
        sr.bold = True
        sr.font.color.rgb = RGBColor(0x1E, 0x7E, 0x34) if approved else RGBColor(0xA0, 0x74, 0x0A)

    for idx, p in enumerate(pages or []):
        if idx > 0:
            doc.add_page_break()
        a = p.get("analysis") or {}
        rot = p.get("rotation_applied") or 0
        rot_lbl = f"  (סובב {rot}°)" if rot else ""
        rtl(doc.add_heading(f"עמוד {p.get('page', idx + 1)}{rot_lbl}", level=2))

        if p.get("image_b64"):
            try:
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(
                    io.BytesIO(base64.b64decode(p["image_b64"])), width=Inches(3.2))
            except Exception:
                log.warning("docx: failed to embed image for page %s", p.get("page"))

        if a.get("page_summary"):
            rtl(doc.add_paragraph(a["page_summary"]))
        if a.get("has_diagram"):
            d = rtl(doc.add_paragraph())
            r = d.add_run(f"▣ סרטוט: {a.get('diagram_description', '')}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x2E, 0x72, 0x86)

        problems = a.get("problems") or []
        if not problems:
            ip = rtl(doc.add_paragraph())
            ip.add_run("לא זוהו תרגילים בעמוד זה.").italic = True

        for pr in problems:
            v = pr.get("verdict", "unclear")
            h = rtl(doc.add_heading(level=3))
            hr = h.add_run(f"תרגיל {pr.get('id', '?')}  ·  {pr.get('topic', '')}  ·  "
                           f"{VERDICT_HE.get(v, v)}")
            hr.font.color.rgb = RGBColor.from_string(VERDICT_COLOR.get(v, "5a6b82"))

            if pr.get("statement_latex"):
                sp = doc.add_paragraph()
                sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sp.add_run("נתון: ").bold = True
                math_run(sp, pr["statement_latex"])

            for s in pr.get("student_steps") or []:
                mark = "✗ " if s.get("ok") is False else ("✓ " if s.get("ok") is True else "• ")
                stp = doc.add_paragraph()
                stp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                mk = stp.add_run(mark)
                mk.bold = True
                mk.font.color.rgb = RGBColor(0xB5, 0x43, 0x43) if s.get("ok") is False \
                    else (RGBColor(0x1E, 0x7E, 0x34) if s.get("ok") is True else RGBColor(0x5A, 0x6B, 0x82))
                math_run(stp, s.get("latex", ""))
                if s.get("comment"):
                    cp = rtl(doc.add_paragraph())
                    cr = cp.add_run(f"⚠ {s['comment']}")
                    cr.italic = True
                    cr.font.size = Pt(9.5)
                    cr.font.color.rgb = RGBColor(0xB5, 0x43, 0x43)

            if pr.get("final_answer_latex"):
                fp = doc.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fp.add_run("תשובה: ").bold = True
                math_run(fp, pr["final_answer_latex"])

            if _num(pr.get("points_max")) is not None:
                scp = rtl(doc.add_paragraph())
                scp.add_run(
                    f"ניקוד: {_fmt_pts(pr.get('points_earned'))} / "
                    f"{_fmt_pts(pr.get('points_max'))}").bold = True
                if pr.get("score_suggestion"):
                    note = scp.add_run(f"   ({pr['score_suggestion']})")
                    note.italic = True
                    note.font.size = Pt(9.5)
                    note.font.color.rgb = RGBColor(0x6E, 0x6A, 0x5C)
            elif pr.get("score_suggestion"):
                rtl(doc.add_paragraph()).add_run(pr["score_suggestion"]).bold = True
            if pr.get("feedback"):
                fb = rtl(doc.add_paragraph())
                fb.add_run(pr["feedback"]).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
